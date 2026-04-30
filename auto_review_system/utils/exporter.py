import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(markdown_text, doc_title="万科智能审计意见函"):
    from docx.oxml.ns import qn
    doc = Document()
    
    # 统一修改文档的默认正文样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 构建红头文件大标题
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(doc_title)
    run.font.name = 'Microsoft YaHei'  # 统一微软雅黑
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(22)
    # Vanke Red
    run.font.color.rgb = RGBColor(229, 0, 18)
    
    # 加个空行副标题区分
    doc.add_paragraph()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. 处理标题 (#, ##, ###)
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            if level > 4: level = 4
            text = m.group(2).strip()
            # 一级标题当正文大标题了，这儿从真实一级开始
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            # 把前3级标题强制设为黑色，显得正式严肃
            if level <= 3:
                run.font.color.rgb = RGBColor(0, 0, 0)
                
        # 2. 处理无序列表项 (- 或 *)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # 兼容特工输出的 **加粗** 特征
            _add_formatted_runs(p, line[2:].strip())
            
        # 3. 处理正常段落
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
            
    # 把 Document 对象保存为 Bytes 流，供 Streamlit 前端下载按钮直通车使用
    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff

def _add_formatted_runs(paragraph, text):
    """
    轻量级 Markdown 解析器引擎
    用于支持 `**粗体**` 的原生 Docx Run 生成解析
    """
    from docx.oxml.ns import qn
    # 以 ** 为界分割字符串，保留匹配项用于判定
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # 掐头去尾
            pure_text = part[2:-2]
            r = paragraph.add_run(pure_text)
            r.bold = True
        else:
            r = paragraph.add_run(part)
            
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
