from docx import Document
import os
import re

def parse_word_doc(file_path):
    """
    (Legacy) 解析 Word 方案，提取纯文本。兼容旧版调用。
    """
    chunks = parse_word_doc_structured(file_path)
    if isinstance(chunks, str):
        return chunks
    return "\\n".join([f"### {c['heading']}\\n{c['text']}" for c in chunks])

def parse_word_as_cost_context(word_path):
    """
    将造价/清单类的 Word 文档榨取为造价上下文 (cost_context)。
    抓取其中所有的文本和尤其是表格数据！
    """
    import docx
    try:
        doc = docx.Document(word_path)
    except Exception as e:
        return f"Error opening Word document: {e}"
        
    context_lines = []
    
    # 提取所有段落
    for p in doc.paragraphs:
        if p.text.strip():
            context_lines.append(p.text.strip())
            
    # 提取所有表格（很多造价清单都在Word表格里）
    for table_idx, table in enumerate(doc.tables, 1):
        context_lines.append(f"--- 造价明细表 {table_idx} ---")
        for row in table.rows:
            row_text = " | ".join([cell.text.strip().replace("\\n", " ") for cell in row.cells if cell.text.strip()])
            if row_text:
                context_lines.append(row_text)
            
    return "\\n".join(context_lines)

def parse_word_doc_structured(file_path):
    """
    (V2.0) 结构化解析 Word 方案，利用 Heading 和正则匹配拆分为独立语义块。
    返回结构: [{'heading': '通用章节', 'text': '段落内容...'}, ...]
    """
    if not os.path.exists(file_path):
        return []
        
    try:
        with open(file_path, 'rb') as f:
            header = f.read(4)
            if header == b'\xd0\xcf\x11\xe0':
                return "【解析异常】: 该文件似乎是旧版的 .doc 格式。本系统仅支持现代 .docx 格式。请将文档另存为 .docx 后重试。"
                
        doc = Document(file_path)
        
        # 章节匹配正则： "一、", "1.", "1.1", "第一章"
        heading_pattern = re.compile(r'^([一二三四五六七八九十]+、|\d+\.\d+|\d+、|第[一二三四五六七八九十]章)')
        
        sections = []
        current_heading = "文档开头(前言)"
        current_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 判断是否为标题层级
            style_name = para.style.name.lower()
            is_heading = 'heading' in style_name or heading_pattern.match(text)
            
            # 避免提取太长的段落作为标题
            if is_heading and len(text) < 40:
                # 碰到新章节，保存旧章节
                if current_text:
                    sections.append({
                        "heading": current_heading,
                        "text": "\\n".join(current_text)
                    })
                current_heading = text
                current_text = []
            else:
                current_text.append(text)
                
        # 保存最后一个章节
        if current_text:
            sections.append({
                "heading": current_heading,
                "text": "\\n".join(current_text)
            })
            
        # 表格数据独立切割为原子切片
        for idx, table in enumerate(doc.tables, 1):
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_text.append(" | ".join(row_data))
            if table_text:
                sections.append({
                    "heading": f"核心数据表 {idx}",
                    "text": "\n".join(table_text)
                })
             
        return sections
        
    except Exception as e:
        return f"【解析异常】: 无法读取该文档内容，内部解析错误 -> {str(e)}"
