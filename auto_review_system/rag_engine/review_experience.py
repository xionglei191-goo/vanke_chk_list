"""
Review-experience extraction and matching for small repair projects.

The source material is expert review history. It is intentionally treated as
experience, not as hard normative text, unless a local standard source can be
identified with confidence.
"""
import datetime as _dt
import difflib
import hashlib
import json
import os
import re
import csv
from collections import Counter, defaultdict
from pathlib import Path

from utils.paths import DATA_DIR, PROJECT_DIR

ANALYSIS_DIR = os.path.join(DATA_DIR, "analysis")
DEFAULT_OPINION_FILE = os.path.join(PROJECT_DIR, "原始材料", "审核意见.xlsx")
DEFAULT_MATERIAL_DIR = os.path.join(PROJECT_DIR, "原始材料", "方案评审")

CORE_DIMENSIONS = ("描述完整性", "工艺合理性", "分项拆分", "逻辑自洽")

WORK_CATEGORY_RULES = [
    ("防水渗漏", ("防水", "渗漏", "渗水", "漏水", "天面", "屋面", "聚氨酯", "JS", "涂膜", "卷材")),
    ("地坪/EPDM/环氧", ("地坪", "环氧", "EPDM", "塑胶", "自流平", "篮球场", "固化", "地垫")),
    ("墙面涂料", ("墙面", "涂料", "油漆", "乳胶漆", "腻子", "真石漆", "外墙翻新")),
    ("瓷砖铺贴", ("瓷砖", "地砖", "墙砖", "玻化砖", "铺贴", "胶泥", "背胶", "湿铺")),
    ("门窗玻璃", ("防火门", "门窗", "玻璃", "钢化玻璃", "雨棚", "3C", "大理石")),
    ("给排水管网", ("管网", "水管", "排水", "排污", "雨水", "污水", "水沟", "水泵", "水箱", "管井")),
    ("弱电设备", ("监控", "摄像", "高空抛物", "人脸识别", "门禁", "交换机", "硬盘", "网线")),
    ("结构修补", ("混凝土", "植筋", "锚固", "钢筋", "反坎", "隔层", "砌筑", "护栏", "钢结构")),
    ("装修翻新", ("宿舍", "卫生间", "天花", "活动室", "装修", "翻新", "改造")),
]

DIMENSION_PATTERNS = {
    "描述完整性": re.compile(
        r"未明确|缺失|缺少|遗漏|需明确|未注明|是什么|多少|哪|是否|规格|型号|厚度|强度|"
        r"配比|比例|参数|尺寸|材质|品牌|范围|部位|做法|类型|品类|验收项|检测"
    ),
    "工艺合理性": re.compile(
        r"不建议|建议|不得|禁止|不应|应采用|需采用|必须|容易|不适合|为什么|选择|"
        r"工艺|湿铺|干铺|胶泥|背胶|聚氨酯|红砖|轻质砂浆|冷涂"
    ),
    "分项拆分": re.compile(
        r"白单|报价|清单|开项|增项|扣除|计量|结算|项目特征|工程量|范围|部位|区分|"
        r"对应|拆除|恢复|界面|按实"
    ),
    "逻辑自洽": re.compile(
        r"复核|核实|确认|不符|不匹配|匹配|冲突|矛盾|逻辑|工序|先|后|再|怎么还有|"
        r"倒置|穿插|干燥|固化|养护|标高|前后"
    ),
}

EVIDENCE_HINTS = [
    ("规范", re.compile(r"不得|必须|严禁|强制|3C|合格证|检测报告|闭水|复验|MU\d+|C\d+")),
    ("方案内部逻辑", re.compile(r"复核|不符|不匹配|冲突|矛盾|怎么还有|先|后|再|标高|对应|型号")),
]

STANDARD_HINTS = [
    ("GB50210-2018 建筑装饰装修工程质量验收标准", ("瓷砖", "涂饰", "抹灰", "门窗", "玻璃", "基层", "腻子")),
    ("GB50208-2011 地下防水工程质量验收规范", ("地下防水", "涂膜", "聚氨酯", "防水涂料", "闭水")),
    ("JGJ/T 112 民用建筑修缮工程施工标准", ("修缮", "植筋", "裂缝", "加固", "既有")),
    ("GB50204-2015 混凝土结构工程施工质量验收规范", ("混凝土", "钢筋", "植筋", "锚固", "浇筑")),
    ("GB50268-2008 给水排水管道工程施工及验收规范", ("给水", "排水", "雨水", "污水", "管道", "沟槽")),
    ("JGJ46 施工现场临时用电安全技术规范", ("临电", "用电", "电工", "漏保", "配电")),
    ("历史审核经验：零星工程专家意见", ()),
]

EXTENSION_RULES = [
    ("EPDM", ("检查胶水配比、固化时间、基层验收、底层/面层厚度、接缝收口和成品保护。", "描述完整性")),
    ("环氧", ("检查基层打磨、裂缝处理、中涂/腻子/面涂顺序、各层厚度、固化保护和画线做法。", "逻辑自洽")),
    ("水沟", ("检查水沟修复与相邻面层的先后顺序、接口顺直、功能测试和成品保护。", "逻辑自洽")),
    ("轻质砖", ("卫生间等有水房间应检查混凝土反坎、砌块强度、砂浆强度和防水交接。", "工艺合理性")),
    ("抹灰", ("检查抹灰厚度、砂浆类型、基层处理、空鼓开裂控制和养护。", "描述完整性")),
    ("植筋", ("检查植筋深度、孔位错开、锚固长度、结构风险和隐蔽验收记录。", "描述完整性")),
    ("聚氨酯", ("检查适用部位；有外饰面或需后续饰面粘结的部位不宜笼统采用聚氨酯。", "工艺合理性")),
    ("瓷砖", ("检查瓷砖规格、吸水率、防滑系数、铺贴方式、胶粘材料、空鼓检查和高低差验收。", "描述完整性")),
    ("钢化玻璃", ("检查 3C 标识、厚度、规格、安装节点、密封收口和破损更换范围。", "描述完整性")),
    ("大理石", ("检查石材规格、六面防护剂、铺贴胶粘材料、空鼓和污染控制。", "描述完整性")),
    ("防火门", ("检查防火等级、五金闭门器、门框灌浆/固定、消防验收和成品保护。", "描述完整性")),
    ("管井", ("检查砌块材料、砂浆强度、井盖材质、承载等级和管井防水/排水。", "工艺合理性")),
    ("摄像", ("检查像素、帧率、宽动态、最低照度、红外补光、存储天数、交换机供电和网线材质。", "描述完整性")),
    ("人脸识别", ("检查设备型号、供电、安装方式、线缆规格、联动调试、试运行和原系统改接责任。", "描述完整性")),
    ("脚手架", ("检查搭设方式、四边防护、抛撑/连墙、验收和高处作业防护。", "工艺合理性")),
]

PROBLEM_PATTERN_RULES = [
    (
        "missing_parameter",
        "关键参数缺失",
        re.compile(r"未明确|需明确|请明确|明确|缺失|缺少|遗漏|是什么|多少|哪|规格|型号|厚度|强度|配比|比例|参数|尺寸|材质|壁厚|间距|范围|部位|品类|容量|转速|等级|标准|功率|电压|分辨率|亮度|拼缝|端口|DN|写错|笔误|混乱|二选一|还是|？|\?"),
        "方案编写停留在工序名称，没有把材料、规格、厚度、强度、配比、部位等关键控制点写成可执行要求。",
        "看到材料或分项工程名称时，反查是否写明规格型号、厚度强度、配比比例、适用部位、施工范围和验收指标。",
    ),
    (
        "method_mismatch",
        "工艺材料不适配",
        re.compile(r"不建议|不得|不应|禁止|建议|应采用|需采用|必须|要求|为什么|选择|湿铺|干铺|聚氨酯|JS|红砖|PVC|冷涂|轻质砂浆|C2TE|不合理|不太合适|风险|防坠|防护|防腐|防锈"),
        "方案套用了通用做法，但没有校核现场环境、基层条件、材料相容性、耐久性或使用场景。",
        "遇到防水、铺贴、涂料、管井、弱电等分项时，检查材料选择是否与基层、环境、振动、潮湿、外饰面或承载条件匹配。",
    ),
    (
        "work_split_or_pricing",
        "分项与计价口径不清",
        re.compile(r"白单|报价|清单|开项|增项|扣除|计量|结算|项目特征|工程量|按实|对应|对下|利旧|恢复|缺项|台班|措施费"),
        "施工动作、报价项目和责任边界没有一一对应，后续容易漏项、重复计价或对下结算争议。",
        "把拆除、基层处理、修复、恢复、检测、成品保护和临时措施逐项对照清单，确认每个动作有对应计价口径。",
    ),
    (
        "sequence_or_logic",
        "工序顺序或逻辑不自洽",
        re.compile(r"复核|不符|不匹配|冲突|矛盾|怎么还有|工序|先|后|再|倒置|交接|固化|养护|保护|标高|前后|逻辑|接线|连接|从下往上|混淆|搞反"),
        "方案没有校核前后工序、交接面、固化养护和现场标高等逻辑，文字存在但落地会返工或失控。",
        "检查工序是否满足先拆除/修复/验收再覆盖、先基层处理再面层、固化养护后开放使用、相邻分项交接先后清楚。",
    ),
    (
        "site_condition_unverified",
        "现场条件未确认",
        re.compile(r"确认|核实|现场|原基层|现有|有无|划痕|旧有|标高|原装饰面|是否需要|需求|甲方|看不出|哪里|几栋|多高|能否|是否|原有|当前"),
        "方案没有把现场调查结果转化为施工依据，导致做法、工程量或修复深度无法判断。",
        "凡涉及既有基层、旧饰面、管线、裂缝、划痕、渗漏路径和甲方需求时，要求写明现场确认结果和对应处理分支。",
    ),
    (
        "acceptance_missing",
        "验收和复核指标缺失",
        re.compile(r"验收|检测|复验|检查|报告|合格证|型式检验|闭水|空鼓|3C|承载|试运行|调试|通电运行"),
        "方案缺少可验收证据或复核动作，无法证明材料合格、施工有效或功能恢复。",
        "对安全玻璃、防火门、防水、瓷砖、弱电、管网等分项，补充证书报告、试验检测、闭水/通水/试运行和现场检查方法。",
    ),
    (
        "interface_scope",
        "界面范围不清",
        re.compile(r"界面|范围|部位|施工内容|移交|甲方|物业|利旧|更新|拆除和恢复|责任|包含|不包含"),
        "方案没有讲清谁做、做哪里、做到什么边界，容易造成审批、施工和结算责任不一致。",
        "逐项明确施工范围、保留/利旧内容、拆除恢复边界、移交条件和不包含事项。",
    ),
]

PATTERN_LABELS = {code: label for code, label, *_ in PROBLEM_PATTERN_RULES}

INTENT_BY_PATTERN = {
    "missing_parameter": ["指导施工", "便于验收", "支撑复核"],
    "method_mismatch": ["指导施工", "控制质量风险"],
    "work_split_or_pricing": ["支撑计价", "避免结算争议"],
    "sequence_or_logic": ["指导施工", "避免返工"],
    "site_condition_unverified": ["支撑复核", "避免错判现场"],
    "acceptance_missing": ["便于验收", "留存复核证据"],
    "interface_scope": ["明确界面", "支撑计价"],
}

PROFESSIONAL_ATTRIBUTION_RULES = [
    {
        "code": "review_object_and_goal",
        "label": "审核对象与维修目标未说清",
        "pattern": re.compile(r"解决的问题|到底是要修啥|需求|目的|阵地修缮标准|项目名称|项目地址|分项工程名称|楼栋号|楼盘名|模板不对|施工内容|不是楼板|外墙开洞"),
        "engineer_question": "这个项目到底要解决哪个现场问题，审批对象是否写对，方案是不是拿错模板或套错场景？",
        "why_it_matters": "零星工程规模小但现场差异大，目标不清会导致后续材料、做法、清单和验收全部跑偏。",
        "transfer_principle": "任何方案先核项目名称、维修对象、现状问题、目标效果和适用模板，目标不清不进入细部审核。",
        "review_questions": ["项目名称和楼栋/部位是否一致", "方案是否说明要解决的病害或功能问题", "是否存在套模板导致的无关内容"],
        "required_artifacts": ["项目名称", "维修对象", "现状问题描述", "目标效果"],
    },
    {
        "code": "as_is_site_basis",
        "label": "现场现状依据不足",
        "pattern": re.compile(r"原基层|旧基层|基层是|基层条件|现场基层|现有|有无|现场|旧有|原装饰面|划痕|水源|源头|哪里|几栋|多高|是否需要|能否|当前|照片|图片|查找|多少|几|是否涉及"),
        "engineer_question": "方案有没有把现场勘查结果写清楚，还是在不知道原状的情况下直接套做法？",
        "why_it_matters": "零星改造是在既有条件上施工，原基层、原饰面、管线、裂缝、渗漏路径不清，做法和工程量都会失真。",
        "transfer_principle": "遇到既有基层、旧饰面、渗漏、裂缝、划痕、管线、标高时，必须先要求现场确认结果和处理分支。",
        "review_questions": ["原基层/原饰面是否写明", "病害原因或水源是否查明", "有无管线/标高/荷载等现场限制", "是否需要按不同现场条件分支处理"],
        "required_artifacts": ["现场照片或描述", "原基层/原材料", "病害原因", "处理分支"],
    },
    {
        "code": "scope_pricing_closure",
        "label": "方案动作与白单/清单未闭环",
        "pattern": re.compile(r"白单|报价|清单|开项|增项|扣除|计量|结算|项目特征|工程量|按实|对下|台班|措施费|缺项|利旧|更新|恢复|不同部位|区分|分别|分部位|区别|为何要分开|有何区别"),
        "engineer_question": "方案里的每个施工动作，在白单/清单里有没有对应项目；清单项目在方案里有没有做法支撑？",
        "why_it_matters": "零星工程最容易在拆除、恢复、措施、利旧和按实台班上发生漏项、重复计价或对下结算争议。",
        "transfer_principle": "把拆除、基层处理、修复、恢复、检测、成品保护、临时措施逐项对照清单，形成方案-计价闭环。",
        "review_questions": ["方案动作是否都有清单项目", "清单项目是否都有做法", "是否存在应扣除或按实结算内容", "利旧/更新边界是否清楚"],
        "required_artifacts": ["白单/清单对应关系", "项目特征", "工程量口径", "措施项目"],
    },
    {
        "code": "material_system_parameters",
        "label": "材料系统性能参数不足",
        "pattern": re.compile(r"规格|型号|厚度|强度|配比|比例|参数|尺寸|材质|壁厚|间距|品类|容量|转速|功率|电压|分辨率|亮度|拼缝|端口|带宽|帧率|刷新率|WDR|照度|等级|标号|DN|304|201|3C|S标|防紫外线|防滑系数|系数|网线|无氧铜|铜包铝|到底|还是|二选一|厚|宽度|深度|种类|品种|成份|成分|粒径|内存|单模|油漆|遍|EPDM|胶水|固化|基层验收"),
        "engineer_question": "班组、采购和验收能不能仅凭方案确定材料性能和安装参数？",
        "why_it_matters": "参数缺失会让班组自由替代材料，造成质量、功能、成本和验收争议。",
        "transfer_principle": "看到任何材料或设备名称，都反查规格、性能、厚度/强度、安装间距、容量功率、认证等级和验收指标。",
        "review_questions": ["材料规格型号是否唯一", "关键性能是否可验收", "设备容量/功率/带宽是否支撑功能", "材料等级是否满足场景"],
        "required_artifacts": ["材料规格表", "设备参数", "认证标志", "允许偏差/验收指标"],
    },
    {
        "code": "method_system_compatibility",
        "label": "工艺系统与场景不匹配",
        "pattern": re.compile(r"不建议|不得|不应|禁止|为什么|选择|应采用|需采用|湿铺|干铺|聚氨酯|JS|红砖|PVC|冷涂|C2TE|XPS|EPS|防水套管|轻质砂浆|透层|粘层|不合理|不太常见|烧焊|卡箍|界面剂|底漆|功能.*重复|疏散方向|防火配件|钢管\\+镀锌"),
        "engineer_question": "这个做法为什么适合当前基层、环境、使用功能和耐久要求？有没有更匹配的材料系统？",
        "why_it_matters": "零星方案常套用通用做法，但不同基层、潮湿、振动、外饰面、承载和消防场景会直接改变工艺选择。",
        "transfer_principle": "材料和工艺不是看能不能做，而是看是否适配基层、环境、功能、耐久和后续饰面系统。",
        "review_questions": ["材料是否适合基层和环境", "是否有振动/潮湿/外饰面/消防要求", "推荐工艺是否写明适用条件", "是否存在禁止或不宜使用材料"],
        "required_artifacts": ["工艺适用条件", "基层处理要求", "替代材料说明", "相容性说明"],
    },
    {
        "code": "detail_node_constructability",
        "label": "构造节点和细部收口不可施工",
        "pattern": re.compile(r"反坎|收口|管根|附加层|V型槽|倒角|美纹纸|滴水线|导流槽|开槽|封堵|空位注浆|洞口|接缝|顺直|节点|侧边|上翻|阴阳角|角条|构造|流槽|沉泥|内外抹灰|什么井"),
        "engineer_question": "关键节点是否写到班组能照着做，收口、搭接、防水、固定和保护有没有落地做法？",
        "why_it_matters": "零星工程失败常不在大工序，而在边角、管根、洞口、接缝、反坎、倒角和新旧交界。",
        "transfer_principle": "凡出现水、缝、洞、边、角、交接、固定和新旧材料连接，都必须审核节点做法。",
        "review_questions": ["节点部位是否识别", "收口/搭接/固定是否明确", "是否有防渗、防裂、防污染措施", "新旧交界是否有处理顺序"],
        "required_artifacts": ["节点做法", "收口措施", "搭接尺寸", "成品保护"],
    },
    {
        "code": "sequence_protection_logic",
        "label": "施工顺序、保护和养护逻辑不成立",
        "pattern": re.compile(r"工序|先|后|再|怎么还有|固化|养护|成品保护|交接|标高|从下往上|混淆|搞反|覆盖|开放使用|顺序|分两层|摊铺|打磨"),
        "engineer_question": "方案的先后顺序、固化养护、成品保护和相邻分项交接是否能避免返工？",
        "why_it_matters": "零星工程往往多工种穿插，顺序错了会导致污染、破坏、标高冲突、返工或责任不清。",
        "transfer_principle": "审核每个面层覆盖前的基层验收、相邻分项交接、固化养护时间和开放使用条件。",
        "review_questions": ["是否先完成隐蔽/基层/功能测试再覆盖", "固化养护和保护时间是否写明", "相邻分项谁先谁后", "标高厚度是否前后一致"],
        "required_artifacts": ["施工流程", "工序间隔", "保护措施", "开放使用条件"],
    },
    {
        "code": "acceptance_records",
        "label": "验收证据和复核动作缺失",
        "pattern": re.compile(r"验收|检测|复验|检查|报告|合格证|型式检验|闭水|空鼓|试运行|调试|通电运行|铭牌|认证|存储天数|施工效果|交付标准|色差|质感"),
        "engineer_question": "工程做完以后，靠什么证据证明材料合格、功能恢复、质量可接受？",
        "why_it_matters": "没有验收动作和资料要求，问题会在结算或质保阶段才暴露，且很难追责。",
        "transfer_principle": "对防水、瓷砖、玻璃、防火门、弱电、管网等分项，必须要求试验、证书、报告或现场检查方法。",
        "review_questions": ["是否有材料合格证明", "是否有功能试验", "是否有现场验收方法", "是否有复核频次和记录"],
        "required_artifacts": ["合格证/检测报告", "试验记录", "验收指标", "影像或记录表"],
    },
    {
        "code": "temporary_safety_access",
        "label": "临时措施、安全通道和施工可达性遗漏",
        "pattern": re.compile(r"脚手架|高空|曲臂车|吊车|吊绳|防坠|临电|防护栏杆|抛撑|围挡|施工时间|工期|备货|台班|物体防坠"),
        "engineer_question": "班组有没有安全、可达、可计价地完成这个零星作业？临时措施是否漏写漏算？",
        "why_it_matters": "零星维修常在高空、通行区域、运营场景中施工，临时措施漏掉会带来安全风险和费用争议。",
        "transfer_principle": "只在高空、临边、吊装、车行、运营中施工等明确触发场景审安全和临措，避免泛化成大而全安全审查。",
        "review_questions": ["是否需要脚手架/吊车/曲臂车/吊绳", "安全防护是否可落地", "临时措施是否进入计价", "工期和备货是否可行"],
        "required_artifacts": ["临时措施", "安全防护", "设备台班", "施工组织约束"],
    },
    {
        "code": "capacity_function_calculation",
        "label": "功能容量和工程量验算不足",
        "pattern": re.compile(r"存储|供电|带宽|解码|水压|排水|坡度|承载|载流量|水泵|井盖|机动车|车道|吨|流量|功耗|30天|72小时|太大|支护"),
        "engineer_question": "参数是否经过功能验算，能不能满足排水、供电、监控存储、承载或耐久需求？",
        "why_it_matters": "设备和土建小改造常被当成材料清单，但实际是否可用取决于容量、荷载、坡度、供电和传输距离。",
        "transfer_principle": "遇到设备系统、管网、道路、井盖、泵、电缆、监控存储时，要求进行容量或功能闭合校核。",
        "review_questions": ["容量是否满足需求", "荷载或承载等级是否匹配", "传输/供电距离是否满足", "坡度/流量/水压是否可实现"],
        "required_artifacts": ["容量计算", "功能参数", "荷载等级", "试运行要求"],
    },
    {
        "code": "expression_drawability",
        "label": "表达方式不足以指导班组施工",
        "pattern": re.compile(r"简图|手绘|线路|逻辑说清楚|整体线路|构件长度|图示|看懂|示意|连接方式"),
        "engineer_question": "文字是否足够让班组理解空间关系、线路逻辑、构件长度和安装位置？是否需要图示？",
        "why_it_matters": "很多零星活靠班组现场理解，复杂节点、线路和骨架如果没有图示，容易做错或报价看不懂。",
        "transfer_principle": "遇到线路、骨架、门禁联动、构件长度、空间关系时，要求补简图或节点示意。",
        "review_questions": ["是否需要节点图/线路图", "构件尺寸和位置是否能看懂", "连接逻辑是否清楚", "班组是否能直接施工"],
        "required_artifacts": ["节点示意", "线路图", "构件尺寸表", "安装位置"],
    },
    {
        "code": "document_data_quality",
        "label": "文本数据质量和名词口径错误",
        "pattern": re.compile(r"名词订正|单位统一|写错|笔误|模板不对|项目名称|分项名称|混乱|ND|DN|隔热层|保温层|干挂法贴地砖"),
        "engineer_question": "方案文字、单位、名词、楼栋和分项是否存在低级错误，是否会误导施工或计价？",
        "why_it_matters": "小项目审批链短，低级文本错误常直接变成采购、施工或结算错误。",
        "transfer_principle": "审核最后必须做一次数据口径清洗：名称、单位、楼栋、分项、材料名词和专业术语统一。",
        "review_questions": ["单位和尺寸是否统一", "专业名词是否正确", "楼栋/部位是否写错", "是否存在模板残留"],
        "required_artifacts": ["统一术语", "单位校核", "楼栋部位校核", "模板清理"],
    },
]

PROFESSIONAL_LABELS = {rule["code"]: rule["label"] for rule in PROFESSIONAL_ATTRIBUTION_RULES}

ALIGNMENT_STATUSES = ("仍缺失", "部分补齐", "已补齐", "无法判断", "无需处理")
CHECKPOINT_STATUSES = ("具体覆盖", "笼统提及", "未覆盖")

SPECIFIC_CHECKPOINT_RULES = [
    ("胶水配比", ("胶水比", "胶水配比", "配比", "比例"), ("胶水比", "胶水配比", "配比", "比例")),
    ("固化/养护时间", ("固化", "养护", "开放使用"), ("固化", "养护", "开放使用", "干燥时间")),
    ("基层验收", ("基层验收", "基层"), ("基层验收", "基层", "平整", "干燥", "含水率", "空鼓")),
    ("水沟交接顺序", ("水沟与EPDM", "水沟交接", "EPDM交接", "先修复水沟"), ("水沟", "接缝", "顺直", "成品保护", "先修复")),
    ("倒角收口", ("倒角", "美纹纸"), ("倒角", "美纹纸", "收口", "遮蔽", "边缘")),
    ("混凝土反坎", ("反坎",), ("反坎", "200mm", "混凝土", "卫生间", "有水房间")),
    ("抹灰厚度", ("抹灰厚度", "10mm", "20mm"), ("抹灰厚度", "10mm", "20mm", "厚度")),
    ("轻质砂浆", ("轻质砂浆",), ("轻质砂浆", "薄抹灰", "砂浆类型")),
    ("植筋深度/锚固", ("植筋", "锚固"), ("植筋", "锚固", "孔径", "孔深", "深度", "拉拔")),
    ("错孔布置", ("错位", "错孔", "水平通缝"), ("错位", "错孔", "水平通缝", "避让")),
    ("结构专业复核", ("后加板", "结构风险", "隔层"), ("后加板", "隔层", "结构复核", "结构风险", "替代方案")),
    ("腻子基层处理", ("腻子", "刮腻子"), ("腻子", "铲除", "打磨", "基层处理", "修补")),
    ("油漆遍数", ("乳胶漆", "1底1面", "一底一面", "底漆", "面漆", "油漆为什么"), ("油漆", "乳胶漆", "底漆", "面漆", "1底1面", "一底一面", "遍")),
    ("角铁规格", ("角铁",), ("角铁", "50×50×5", "50*50*5", "规格", "尺寸")),
    ("方通使用部位", ("方通", "方管"), ("方通", "方管", "使用部位", "侧边", "骨架", "承重")),
    ("安装间距", ("间距",), ("间距", "400mm", "500mm", "不超过")),
    ("焊接防腐", ("焊接", "防腐", "防锈"), ("焊接", "焊渣", "防腐", "防锈", "底漆", "面漆")),
    ("C2TE性能等级", ("C2TE",), ("C2TE",)),
    ("专用瓷砖胶/粘结剂", ("瓷砖胶", "专用瓷砖胶"), ("瓷砖胶", "专用粘结剂", "胶粘剂", "粘结剂")),
    ("禁止干铺", ("干铺",), ("不得干铺", "禁止干铺", "干铺", "湿铺")),
    ("六面防护剂", ("六面", "防护剂"), ("六面", "防护剂", "石材防护")),
    ("现场滴水检查", ("滴水",), ("滴水", "水珠", "吸收", "滚落", "现场检查")),
    ("3C标识", ("3C",), ("3C", "CCC")),
    ("防火门型式资料", ("型式认可证书", "型式检验报告", "核心资料三查"), ("型式认可证书", "型式检验报告", "检测报告", "合格证")),
    ("产品铭牌", ("产品铭牌", "铭牌"), ("产品铭牌", "铭牌")),
    ("填充密实检查", ("敲击门板", "填充是否密实", "偷工减料"), ("敲击", "门板", "填充", "密实")),
    ("防火门顺序器", ("顺序器", "双开门", "双扇门"), ("顺序器", "双开门", "双扇门", "关闭顺序")),
    ("分项开项", ("分项工程开项", "单独开项", "开项", "注明部位", "备注列中附图"), ("分项工程", "单独开项", "部位", "备注", "附图")),
    ("泄水管", ("泄水管", "排/泄水管", "水压力", "水土流失"), ("泄水管", "排水管", "水压力", "水土流失")),
    ("交换机型号", ("主交换机", "交换机", "型号"), ("交换机", "型号")),
    ("端口分配", ("端口分配",), ("端口", "分配")),
    ("连续通电测试", ("连续通电", "72小时", "通电运行测试"), ("连续通电", "72小时", "通电运行", "测试")),
    ("电箱/立杆规格", ("电箱壁厚", "不锈钢等级", "立杆长宽厚", "立杆", "电箱"), ("电箱", "壁厚", "不锈钢", "立杆", "长宽厚")),
    ("外墙裂缝材料体系", ("快干水泥", "防水抗裂砂浆", "丙烯酸防水涂料", "外墙裂缝"), ("快干水泥", "防水抗裂砂浆", "丙烯酸", "JS", "聚氨酯", "注浆")),
    ("路面划线", ("路面划线", "划线"), ("路面划线", "标线", "划线")),
    ("项目名称校核", ("项目名称填写错误", "项目名称"), ("项目名称", "工程名称")),
    ("工期可控性", ("工期", "备货", "施工时间"), ("工期", "备货", "施工时间", "进度")),
]

GENERIC_ARTIFACT_ALIASES = {
    "项目名称": ("项目名称", "工程名称"),
    "维修对象": ("维修对象", "施工范围", "施工内容", "部位"),
    "现状问题描述": ("现状", "问题", "病害", "渗漏", "破损", "原有"),
    "目标效果": ("目标", "效果", "交付标准", "验收"),
    "现场照片或描述": ("现场", "照片", "现状", "原有"),
    "原基层/原材料": ("原基层", "旧基层", "基层", "原材料", "原装饰面"),
    "病害原因": ("原因", "水源", "渗漏路径", "裂缝"),
    "处理分支": ("如", "若", "分别", "处理分支", "不同情况"),
    "白单/清单对应关系": ("白单", "清单", "报价", "对应"),
    "项目特征": ("项目特征", "特征描述"),
    "工程量口径": ("工程量", "计量", "按实", "扣除"),
    "措施项目": ("措施", "脚手架", "围蔽", "吊车", "台班"),
    "材料规格表": ("规格", "型号", "厚度", "强度", "材质", "等级", "尺寸"),
    "设备参数": ("参数", "功率", "容量", "分辨率", "带宽", "存储"),
    "认证标志": ("认证", "标识", "3C", "S标", "铭牌"),
    "允许偏差/验收指标": ("允许偏差", "验收指标", "偏差", "空鼓", "平整度"),
    "工艺适用条件": ("适用", "基层", "环境", "潮湿", "振动"),
    "基层处理要求": ("基层处理", "基层清理", "打磨", "找平"),
    "替代材料说明": ("替代", "改用", "建议采用"),
    "相容性说明": ("相容", "粘结", "配套", "系统"),
    "节点做法": ("节点", "收口", "附加层", "管根", "洞口"),
    "收口措施": ("收口", "密封", "打胶", "美纹纸"),
    "搭接尺寸": ("搭接", "上翻", "宽度", "尺寸"),
    "成品保护": ("成品保护", "保护", "遮蔽"),
    "施工流程": ("施工流程", "施工工序", "顺序", "先", "后"),
    "工序间隔": ("间隔", "固化", "养护", "干燥"),
    "保护措施": ("保护", "围蔽", "遮挡", "覆盖"),
    "开放使用条件": ("开放", "投入使用", "养护", "固化"),
    "合格证/检测报告": ("合格证", "检测报告", "材质证明", "型式检验"),
    "试验记录": ("试验", "闭水", "通水", "拉拔", "试运行"),
    "验收指标": ("验收", "指标", "偏差", "质量要求"),
    "影像或记录表": ("记录", "影像", "照片", "验收表"),
    "临时措施": ("临时", "脚手架", "吊车", "曲臂车", "围挡"),
    "安全防护": ("防护", "临边", "高处", "防坠"),
    "设备台班": ("台班", "吊车", "机械"),
    "施工组织约束": ("施工时间", "工期", "运营", "通行"),
    "容量计算": ("计算", "容量", "承载", "水压", "流量"),
    "功能参数": ("功能", "参数", "存储", "供电", "带宽"),
    "荷载等级": ("荷载", "承载", "等级", "吨"),
    "试运行要求": ("试运行", "调试", "通电运行"),
    "节点示意": ("节点示意", "简图", "图示"),
    "线路图": ("线路图", "线路", "接线"),
    "构件尺寸表": ("构件", "长度", "尺寸"),
    "安装位置": ("安装位置", "部位", "点位"),
    "统一术语": ("术语", "名称", "名词"),
    "单位校核": ("单位", "mm", "m2", "m³"),
    "楼栋部位校核": ("楼栋", "部位", "栋", "层"),
    "模板清理": ("模板", "保修", "移交状态"),
    "方案原文证据": ("施工范围", "施工工序", "验收", "材料"),
    "专家复核备注": ("复核", "备注", "意见"),
}


def _clean_text(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _hash_id(prefix, *parts):
    h = hashlib.sha256("\n".join(str(p or "") for p in parts).encode("utf-8")).hexdigest()[:12].upper()
    return f"{prefix}_{h}"


def split_opinion_items(opinion):
    text = str(opinion or "").strip()
    if not text or text in {"/", "无", "无意见"}:
        return []
    normalized = text.replace("\r", "\n").replace("\n", " / ")
    parts = re.split(r"(?:^|[；;/]\s*)\d+[、.．]\s*", normalized)
    items = [p.strip(" ；;/") for p in parts if p.strip(" ；;/")]
    return items or [normalized]


def classify_project_type(project_name):
    name = str(project_name or "")
    has_scheme = "施工方案" in name or "工程方案" in name or name.endswith("方案")
    has_cost = "报价" in name or "清单" in name or "白单" in name
    if has_scheme and has_cost:
        return "方案兼报价"
    if has_scheme:
        return "施工方案"
    if has_cost:
        return "报价/白单"
    return "其他资料"


def infer_work_category(text):
    haystack = str(text or "")
    scores = []
    for category, keywords in WORK_CATEGORY_RULES:
        score = sum(1 for keyword in keywords if keyword.lower() in haystack.lower())
        if score:
            scores.append((score, category))
    if not scores:
        return "通用零星维修"
    scores.sort(reverse=True)
    return scores[0][1]


def classify_dimension(opinion):
    text = str(opinion or "")
    scores = []
    for dimension, pattern in DIMENSION_PATTERNS.items():
        hits = pattern.findall(text)
        if hits:
            scores.append((len(hits), dimension))
    if not scores:
        return "描述完整性"
    priority = {name: idx for idx, name in enumerate(CORE_DIMENSIONS)}
    scores.sort(key=lambda item: (-item[0], priority.get(item[1], 99)))
    return scores[0][1]


def classify_problem_patterns(opinion):
    text = str(opinion or "")
    matches = []
    for code, label, pattern, root_cause, generalization_rule in PROBLEM_PATTERN_RULES:
        hits = pattern.findall(text)
        if hits:
            matches.append({
                "code": code,
                "label": label,
                "score": len(hits),
                "root_cause": root_cause,
                "generalization_rule": generalization_rule,
            })
    if not matches:
        matches.append({
            "code": "expert_judgement",
            "label": "专家经验判断",
            "score": 1,
            "root_cause": "意见依赖专家对零星工程场景的经验判断，需要结合分项、现场条件和方案上下文复核。",
            "generalization_rule": "遇到相似分项时，先判断该意见是在补参数、校核做法、拆分计价还是排查逻辑，再迁移到当前方案。",
        })
    matches.sort(key=lambda item: item["score"], reverse=True)
    return matches


def primary_problem_pattern(opinion):
    return classify_problem_patterns(opinion)[0]


def classify_professional_attributions(opinion):
    text = str(opinion or "")
    matches = []
    for idx, rule in enumerate(PROFESSIONAL_ATTRIBUTION_RULES):
        hits = rule["pattern"].findall(text)
        if hits:
            item = {key: value for key, value in rule.items() if key != "pattern"}
            item["score"] = len(hits)
            item["_priority"] = idx
            matches.append(item)
    if not matches:
        fallback = {
            "code": "expert_engineering_judgement",
            "label": "需专业工程师结合上下文判断",
            "engineer_question": "这条意见依赖专家对项目上下文的经验判断，需要回到方案原文确认触发条件。",
            "why_it_matters": "历史意见不能脱离项目照搬，必须先判断当前方案是否出现同类工程对象、现场条件或风险路径。",
            "transfer_principle": "先回到方案原文找触发证据，再决定是否按同类经验迁移。",
            "review_questions": ["当前方案是否出现同类对象", "是否存在相同现场条件", "是否能找到原文证据"],
            "required_artifacts": ["方案原文证据", "专家复核备注"],
            "score": 1,
            "_priority": 999,
        }
        matches.append(fallback)
    matches.sort(key=lambda item: (-item["score"], item["_priority"]))
    for item in matches:
        item.pop("_priority", None)
    return matches


def primary_professional_attribution(opinion):
    return classify_professional_attributions(opinion)[0]


def infer_review_intents(problem_patterns, dimension):
    intents = []
    for pattern in problem_patterns:
        for intent in INTENT_BY_PATTERN.get(pattern["code"], []):
            if intent not in intents:
                intents.append(intent)
    if not intents:
        fallback = {
            "描述完整性": "指导施工",
            "工艺合理性": "控制质量风险",
            "分项拆分": "支撑计价",
            "逻辑自洽": "支撑复核",
        }
        intents.append(fallback.get(dimension, "支撑复核"))
    return intents


def build_attribution(opinion, dimension, work_category):
    patterns = classify_problem_patterns(opinion)
    primary = patterns[0]
    professional = classify_professional_attributions(opinion)
    primary_professional = professional[0]
    return {
        "problem_pattern": primary["code"],
        "problem_pattern_label": primary["label"],
        "secondary_patterns": [item["code"] for item in patterns[1:4]],
        "professional_attribution": primary_professional["code"],
        "professional_attribution_label": primary_professional["label"],
        "secondary_professional_attributions": [item["code"] for item in professional[1:4]],
        "engineer_question": primary_professional["engineer_question"],
        "review_intents": infer_review_intents(patterns, dimension),
        "root_cause": primary_professional["why_it_matters"],
        "risk_if_ignored": infer_risk_if_ignored(opinion, dimension, work_category, primary["code"]),
        "generalization_rule": primary_professional["transfer_principle"],
        "review_questions": primary_professional.get("review_questions", []),
        "required_artifacts": primary_professional.get("required_artifacts", []),
        "fix_template": build_fix_template(opinion, dimension, work_category, primary["code"]),
        "standard_query": build_standard_query(opinion, work_category),
    }


def infer_evidence_type(opinion, dimension):
    text = str(opinion or "")
    for evidence_type, pattern in EVIDENCE_HINTS:
        if pattern.search(text):
            return evidence_type
    if dimension == "逻辑自洽":
        return "方案内部逻辑"
    return "专家经验"


def infer_evidence_ref(opinion, work_category, evidence_type):
    if evidence_type == "方案内部逻辑":
        return "方案文本内部一致性校核"
    if evidence_type != "规范":
        return "历史审核经验：零星工程专家意见"
    text = f"{opinion} {work_category}"
    for ref, keywords in STANDARD_HINTS:
        if keywords and any(keyword.lower() in text.lower() for keyword in keywords):
            return ref
    return "历史审核经验：零星工程专家意见"


def infer_risk_if_ignored(opinion, dimension, work_category, pattern_code):
    text = f"{opinion} {work_category}"
    if "EPDM" in text or "环氧" in text:
        return "可能导致面层起鼓、脱层、开裂、接缝不顺或开放使用时间争议。"
    if "防水" in text or "渗漏" in text:
        return "可能导致渗漏复发、饰面粘结失败、部位漏项或责任难以界定。"
    if "植筋" in text or "混凝土结构" in text or "后加板" in text:
        return "可能形成结构连接隐患、原结构损伤或隐蔽验收无法追溯。"
    if "玻璃" in text or "防火门" in text:
        return "可能使用不合格材料或缺少消防/安全复核凭证。"
    if "瓷砖" in text or "石材" in text or "大理石" in text:
        return "可能出现空鼓、返碱、污染、开裂或高低差等质量问题。"
    if pattern_code == "work_split_or_pricing":
        return "可能造成方案动作无报价、清单项目无做法、重复计价或对下结算争议。"
    if pattern_code == "sequence_or_logic":
        return "可能导致工序返工、成品破坏、参数不匹配或审批后仍无法施工。"
    if dimension == "描述完整性":
        return "班组会按经验自由发挥，现场、验收和复核口径容易失控。"
    if dimension == "工艺合理性":
        return "可能采用不适配做法，影响耐久性、观感或功能恢复。"
    return "可能导致施工、计价、验收和复核之间无法闭环。"


def build_fix_template(opinion, dimension, work_category, pattern_code):
    if pattern_code == "missing_parameter":
        return f"在{work_category}对应章节补充材料规格、型号、厚度/强度、配比、施工部位、验收指标和允许偏差。"
    if pattern_code == "method_mismatch":
        return f"说明{work_category}选择该工艺的适用条件；如与现场不匹配，改用专家意见建议的材料或做法，并写明基层处理。"
    if pattern_code == "work_split_or_pricing":
        return "把方案动作拆成拆除、基层处理、修复、恢复、检测、成品保护等项目，并与白单/清单逐项对应。"
    if pattern_code == "sequence_or_logic":
        return "重排施工流程，明确前置验收、工序间隔、固化/养护时间、相邻分项交接和成品保护。"
    if pattern_code == "site_condition_unverified":
        return "补充现场确认结果；对不同现场条件设置处理分支，并说明工程量、材料和做法如何随条件调整。"
    if pattern_code == "acceptance_missing":
        return "补充合格证/检测报告/试验记录/现场检查方法，并明确验收频次、指标和责任人。"
    if pattern_code == "interface_scope":
        return "明确施工范围、利旧/更新内容、拆除恢复边界、移交条件和不包含事项。"
    return f"结合{work_category}补充问题背景、处理原则、施工参数、验收方法和报价口径。"


def build_standard_query(opinion, work_category):
    parts = []
    for keyword in trigger_keywords_for(opinion, work_category=work_category)[:5]:
        parts.append(keyword)
    if work_category != "通用零星维修":
        parts.append(work_category)
    dimension = classify_dimension(opinion)
    if dimension == "描述完整性":
        parts.extend(["材料参数", "验收"])
    elif dimension == "工艺合理性":
        parts.extend(["施工工艺", "质量控制"])
    elif dimension == "分项拆分":
        parts.extend(["工程量", "项目特征"])
    else:
        parts.extend(["施工顺序", "验收"])
    return " ".join(dict.fromkeys(parts))


def extension_rules_for(opinion, project_name=""):
    text = f"{project_name} {opinion}"
    rules = []
    for keyword, (rule, dimension) in EXTENSION_RULES:
        if keyword.lower() in text.lower():
            rules.append({"trigger": keyword, "rule": rule, "dimension": dimension})
    if not rules:
        rules.append({
            "trigger": infer_work_category(text),
            "rule": "检查材料、工艺、工序、验收、界面和报价口径是否能支撑现场施工与复核。",
            "dimension": classify_dimension(opinion),
        })
    return rules


def trigger_keywords_for(opinion, project_name="", work_category=""):
    text = f"{project_name} {opinion} {work_category}"
    keywords = []
    for category, category_keywords in WORK_CATEGORY_RULES:
        for keyword in category_keywords:
            if keyword.lower() in text.lower() and keyword not in keywords:
                keywords.append(keyword)
    for keyword, _ in EXTENSION_RULES:
        if keyword.lower() in text.lower() and keyword not in keywords:
            keywords.append(keyword)
    for token in re.findall(r"[A-Za-z0-9]{2,}|[\u4e00-\u9fff]{2,8}", text):
        if token in {"广州", "万科", "工程", "施工", "方案", "报价", "清单", "明确", "缺失"}:
            continue
        if len(keywords) >= 10:
            break
        if token not in keywords:
            keywords.append(token)
    return keywords[:10]


def _contains_any(text, aliases):
    haystack = str(text or "").lower()
    return any(str(alias).lower() in haystack for alias in aliases if str(alias).strip())


def _has_numeric_or_threshold(text):
    return bool(re.search(r"\d+\s*(?:[:：]\s*\d+|小时|h|H|天|日|d|D|mm|cm|m|%|遍|级)|[≥≤]|不超过|不小于|不少于", str(text or "")))


def _checkpoint_status(checkpoint, evidence_text):
    name = checkpoint["name"]
    aliases = checkpoint["aliases"]
    has_mention = _contains_any(evidence_text, aliases)
    text = str(evidence_text or "")

    if name == "胶水配比":
        if not has_mention:
            return "未覆盖", "未看到胶水比例或配比要求。"
        exact = re.search(r"(胶水比|胶水配比|配比|比例)[^。；，,]{0,30}(\d+\s*[:：]\s*\d+|\d+\s*(?:%|份)|按[^。；，,]{1,30}比例)", text, re.I)
        if exact:
            return "具体覆盖", "已写出胶水配比或可执行的配比口径。"
        return "笼统提及", "仅提到配比/配比证明，未写明胶水比例或施工配合比。"

    if name == "固化/养护时间":
        if not has_mention:
            return "未覆盖", "未看到固化、养护或开放使用时间。"
        exact = re.search(r"(固化|养护|干燥|开放使用)[^。；，,]{0,40}(\d+\s*(?:小时|h|H|天|日|d|D))|(\d+\s*(?:小时|h|H|天|日|d|D))[^。；，,]{0,40}(固化|养护|干燥|开放使用)", text)
        if exact:
            return "具体覆盖", "已写出固化/养护/开放使用的时间条件。"
        return "笼统提及", "仅写固化或养护动作，未写时间或开放使用条件。"

    if name == "基层验收":
        if not has_mention:
            return "未覆盖", "未看到基层验收或基层条件要求。"
        exact = re.search(r"(基层|地面验收)[^。；]{0,120}(含水率|平整度|强度|空鼓|起砂|干燥|清洁|无杂物|无松动|2m|≤|≥)", text)
        if exact:
            return "具体覆盖", "已把基层验收和具体指标或状态要求关联起来。"
        return "笼统提及", "仅提到地面验收/基层处理，未写基层验收指标。"

    if name == "水沟交接顺序":
        exact = re.search(r"(先[^。；]{0,40}水沟[^。；]{0,80}(成品保护|保护)?[^。；]{0,80}(再|后)[^。；]{0,60}EPDM|水沟[^。；]{0,80}成品保护[^。；]{0,80}EPDM)", text, re.I)
        if exact:
            return "具体覆盖", "已明确水沟先行、保护和 EPDM 后续铺装关系。"
        if re.search(r"EPDM[^。；]{0,220}水沟", text, re.I):
            return "未覆盖", "文本顺序更像先 EPDM 后水沟，未体现专家要求的交接顺序。"
        if has_mention:
            return "笼统提及", "提到水沟或 EPDM，但未写清两者交接先后和保护。"
        return "未覆盖", "未看到水沟与 EPDM 交接控制。"

    if name in {"角铁规格", "方通使用部位", "安装间距", "焊接防腐", "混凝土反坎", "抹灰厚度"}:
        if not has_mention:
            return "未覆盖", f"未看到{name}。"
        if name in {"角铁规格", "安装间距", "混凝土反坎", "抹灰厚度"} and not _has_numeric_or_threshold(text):
            return "笼统提及", f"提到{name}，但未写具体尺寸、厚度或阈值。"
        return "具体覆盖", f"已写明{name}。"

    if name == "油漆遍数":
        if not has_mention:
            return "未覆盖", "未看到底漆/面漆遍数。"
        if re.search(r"(为什么|为何)", str(checkpoint.get("source_opinion", ""))):
            rationale = re.search(
                r"(油漆|乳胶漆|底漆|面漆)[^。；]{0,120}(耐擦洗|遮盖|公共|使用频率|观感|原因|满足)",
                text,
            )
            if not rationale:
                return "笼统提及", "写了遍数，但未说明 1 底 1 面为什么满足空间耐久和观感要求。"
        return "具体覆盖", "已写明底漆和面漆遍数。"

    direct_presence_checkpoints = {
        "C2TE性能等级", "六面防护剂", "现场滴水检查", "3C标识", "轻质砂浆",
        "植筋深度/锚固", "错孔布置", "结构专业复核", "防火门型式资料",
        "产品铭牌", "防火门顺序器", "泄水管", "交换机型号", "端口分配",
        "连续通电测试", "路面划线",
    }
    if name in direct_presence_checkpoints:
        if has_mention:
            return "具体覆盖", f"已写明{name}。"
        return "未覆盖", f"未看到{name}。"

    if name == "填充密实检查":
        if has_mention:
            return "具体覆盖", "已写出门板敲击或填充密实检查。"
        return "未覆盖", "未看到门板填充密实的现场复核方法。"

    if name == "分项开项":
        if not has_mention:
            return "未覆盖", "未看到按不同部位/做法单独开项或备注附图。"
        if _contains_any(text, ("部位", "备注", "附图", "单独开项")):
            return "具体覆盖", "已体现部位、备注或附图等分项开项口径。"
        return "笼统提及", "提到分项，但未写清部位、备注或附图口径。"

    if name == "电箱/立杆规格":
        if not has_mention:
            return "未覆盖", "未看到电箱壁厚、不锈钢等级或立杆长宽厚。"
        if _has_numeric_or_threshold(text) and _contains_any(text, ("壁厚", "不锈钢", "立杆", "电箱")):
            return "具体覆盖", "已写出电箱/立杆规格参数。"
        return "笼统提及", "提到电箱或立杆，但未写明壁厚、材质等级或长宽厚。"

    if name == "外墙裂缝材料体系":
        if _contains_any(text, ("防水抗裂砂浆", "丙烯酸", "JS", "聚氨酯", "注浆")):
            return "具体覆盖", "已写明外墙裂缝或渗漏修补材料体系。"
        if _contains_any(text, ("外墙裂缝", "裂缝", "防水涂料", "快干水泥")):
            return "笼统提及", "提到裂缝/防水材料，但未明确适配材料体系。"
        return "未覆盖", "未看到外墙裂缝修补材料体系。"

    if name == "项目名称校核":
        if has_mention:
            return "笼统提及", "文本包含项目/工程名称，但需人工核对是否填写正确。"
        return "未覆盖", "未看到项目名称或工程名称，无法核对名称是否正确。"

    if name == "工期可控性":
        if has_mention:
            return "笼统提及", "提到工期或施工时间，但仍需复核备货和实际施工可控性。"
        return "未覆盖", "未看到工期、备货或施工时间安排。"

    if not has_mention:
        return "未覆盖", f"未看到{name}。"
    if _has_numeric_or_threshold(text):
        return "具体覆盖", f"已写出{name}的具体参数或验收阈值。"
    return "笼统提及", f"仅提到{name}，还不足以直接施工或验收。"


def _dedupe_checkpoints(checkpoints):
    deduped = []
    seen = set()
    for checkpoint in checkpoints:
        name = checkpoint["name"]
        if name in seen:
            continue
        seen.add(name)
        deduped.append(checkpoint)
    return deduped


def expected_checkpoints_for(opinion, work_category="", attribution=None):
    """Return concrete control points the expert opinion is really asking us to verify."""
    text = f"{opinion} {work_category}"
    checkpoints = []
    for name, triggers, aliases in SPECIFIC_CHECKPOINT_RULES:
        if _contains_any(text, triggers):
            checkpoints.append({"name": name, "aliases": aliases, "source": "opinion", "source_opinion": str(opinion or "")})

    if checkpoints:
        return _dedupe_checkpoints(checkpoints)

    attribution = attribution or {}
    for artifact in attribution.get("required_artifacts", [])[:4]:
        aliases = GENERIC_ARTIFACT_ALIASES.get(artifact, (artifact,))
        checkpoints.append({"name": artifact, "aliases": aliases, "source": "required_artifact", "source_opinion": str(opinion or "")})
    return _dedupe_checkpoints(checkpoints)


def _snippet_around_keywords(text, keywords, max_chars=600):
    text = _clean_text(text)
    if len(text) <= max_chars:
        return text
    lowered = text.lower()
    positions = [
        lowered.find(str(keyword).lower())
        for keyword in keywords
        if str(keyword).strip() and lowered.find(str(keyword).lower()) >= 0
    ]
    if not positions:
        return text[:max_chars]
    center = min(positions)
    start = max(0, center - max_chars // 3)
    end = min(len(text), start + max_chars)
    if end - start < max_chars:
        start = max(0, end - max_chars)
    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end]}{suffix}"


def build_expert_intent(row):
    opinion = row.get("opinion") or row.get("source_opinion", "")
    question = row.get("engineer_question") or primary_professional_attribution(opinion).get("engineer_question", "")
    return f"专家在追问：{question}"


def is_non_actionable_opinion(opinion):
    text = _clean_text(opinion)
    if text in {"同意", "无", "/", ""}:
        return True
    if text.startswith("同意"):
        return True
    if re.fullmatch(r"同意[。.!！]?", text):
        return True
    if text.startswith("亮点"):
        return True
    return False


def assess_scheme_alignment(row):
    """Compare an expert opinion with extracted source-scheme evidence.

    This is intentionally conservative. It does not try to prove compliance;
    it records whether the current source material appears to contain the
    control points implied by the expert's opinion.
    """
    evidence = row.get("scheme_evidence", []) or []
    opinion = row.get("opinion") or row.get("source_opinion", "")
    if is_non_actionable_opinion(opinion):
        return {
            "alignment_status": "无需处理",
            "covered_points": [],
            "partial_points": [],
            "missing_points": [],
            "checkpoint_assessments": [],
            "scheme_gap": "该条为同意或正向评价，不作为缺陷经验卡输出。",
            "expert_intent": "该条不是问题意见，无需转化为缺陷检查规则。",
            "evidence_chain": {
                "source_opinion": opinion,
                "scheme_evidence": evidence[:3],
                "covered_points": [],
                "partial_points": [],
                "missing_points": [],
                "checkpoint_assessments": [],
                "alignment_status": "无需处理",
                "scheme_gap": "该条为同意或正向评价，不作为缺陷经验卡输出。",
                "expert_intent": "该条不是问题意见，无需转化为缺陷检查规则。",
                "transfer_principle": "",
            },
        }
    attribution = {
        "required_artifacts": row.get("required_artifacts", []),
        "review_questions": row.get("review_questions", []),
    }
    checkpoints = expected_checkpoints_for(opinion, row.get("work_category", ""), attribution)
    alignment_text = row.get("alignment_text", "")
    evidence_text = alignment_text or " ".join(item.get("text", "") for item in evidence)
    covered = []
    partial = []
    missing = []
    checkpoint_assessments = []
    for checkpoint in checkpoints:
        checkpoint_status, note = _checkpoint_status(checkpoint, evidence_text)
        checkpoint_assessments.append({
            "name": checkpoint["name"],
            "status": checkpoint_status,
            "note": note,
            "source": checkpoint.get("source", ""),
        })
        if checkpoint_status == "具体覆盖":
            covered.append(checkpoint["name"])
        elif checkpoint_status == "笼统提及":
            partial.append(checkpoint["name"])
        else:
            missing.append(checkpoint["name"])

    if not evidence and not alignment_text:
        status = "无法判断"
        scheme_gap = "未找到可对照的原方案片段，需要先回到方案原文确认该意见的触发位置。"
    elif not checkpoints:
        status = "无法判断"
        scheme_gap = "该意见缺少可机械拆解的控制点，需工程师结合原方案上下文判断。"
    elif covered and not partial and not missing:
        status = "已补齐"
        scheme_gap = "当前对照到的方案证据已覆盖历史意见中的主要控制点；运行时应识别为已补齐或历史版本已修订，不能机械复述为缺陷。"
    elif covered or partial:
        status = "部分补齐"
        parts = []
        if covered:
            parts.append(f"已具体覆盖 {'、'.join(covered)}")
        if partial:
            parts.append(f"仅笼统提及 {'、'.join(partial)}")
        if missing:
            parts.append(f"仍缺失 {'、'.join(missing)}")
        scheme_gap = "；".join(parts) + "。"
    else:
        status = "仍缺失"
        if alignment_text and not evidence:
            scheme_gap = f"已对高可信匹配文件做全文检索，仍未覆盖 {'、'.join(missing)}，历史意见指向的关键控制点仍需补充。"
        else:
            scheme_gap = f"当前方案片段未覆盖 {'、'.join(missing)}，历史意见指向的关键控制点仍需补充。"

    return {
        "alignment_status": status,
        "covered_points": covered,
        "partial_points": partial,
        "missing_points": missing,
        "checkpoint_assessments": checkpoint_assessments,
        "scheme_gap": scheme_gap,
        "expert_intent": build_expert_intent(row),
        "evidence_chain": {
            "source_opinion": opinion,
            "scheme_evidence": evidence[:3],
            "covered_points": covered,
            "partial_points": partial,
            "missing_points": missing,
            "checkpoint_assessments": checkpoint_assessments,
            "alignment_status": status,
            "scheme_gap": scheme_gap,
            "expert_intent": build_expert_intent(row),
            "transfer_principle": row.get("generalization_rule", ""),
        },
    }


def build_background(opinion, dimension, work_category):
    if dimension == "描述完整性":
        return f"{work_category}通常依赖材料规格、做法参数和验收指标来落地，缺项会导致班组按经验施工、甲方无法复核。"
    if dimension == "工艺合理性":
        return f"{work_category}的做法选择会影响耐久性、维修效果和后续饰面/使用功能，不能只写笼统工序。"
    if dimension == "分项拆分":
        return f"{work_category}涉及多个部位或工序时，方案、清单和责任界面需要拆清，否则容易漏项、重复计价或结算扯皮。"
    return f"{work_category}需要前后工序、参数和现场条件相互匹配，否则即使文字齐全也无法指导施工。"


def build_reason(opinion, dimension):
    text = _clean_text(opinion)
    if dimension == "描述完整性":
        return f"专家关注点不是文字多少，而是关键参数是否可验收、可采购、可结算。该意见指出“{text}”，说明方案缺少可执行控制点。"
    if dimension == "工艺合理性":
        return f"该意见指出“{text}”，核心是在校核做法是否适配现场、材料和使用场景，避免采用不耐久或不匹配的工艺。"
    if dimension == "分项拆分":
        return f"该意见指出“{text}”，核心是把施工动作、报价口径和责任边界拆清，防止后续签证或对下结算争议。"
    return f"该意见指出“{text}”，核心是发现方案内部的顺序、参数或对象不一致，要求先复核逻辑再审批。"


def index_material_files(material_dir=DEFAULT_MATERIAL_DIR):
    base = Path(material_dir)
    if not base.exists():
        return []
    return [p for p in base.iterdir() if p.is_file() and not p.name.startswith(".~")]


GENERIC_NAME_WORDS = {
    "广州", "万科", "工程", "报价", "清单", "施工", "方案", "项目", "维修", "改造",
    "翻新", "零星", "花园", "小区", "专项", "二类", "一类",
}


def _normalized_material_name(value):
    text = _clean_text(value)
    text = re.sub(r"F\d{2}", "", text, flags=re.I)
    text = re.sub(r"\d{6,8}", "", text)
    text = re.sub(r"[（(].*?[）)]", "", text)
    for word in GENERIC_NAME_WORDS:
        text = text.replace(word, "")
    text = re.sub(r"[\s_+\-—、，,：:（）()]+", "", text)
    return text


def material_match_quality(project_name, material_file):
    if not material_file:
        return {"label": "未匹配", "score": 0.0}
    project = _normalized_material_name(project_name)
    matched = _normalized_material_name(Path(material_file).stem)
    if not project or not matched:
        return {"label": "需人工确认", "score": 0.0}
    score = difflib.SequenceMatcher(None, project, matched).ratio()
    if project in matched or matched in project:
        label = "高"
    elif score >= 0.68:
        label = "中"
    else:
        label = "低，疑似错配"
    return {"label": label, "score": round(score, 3)}


def match_material_file(project_name, files):
    if not files:
        return None
    stems = [p.stem for p in files]
    cleaned = str(project_name or "").replace("报价清单", "").replace("施工方案", "").replace("工程方案", "").strip()
    matches = difflib.get_close_matches(cleaned, stems, n=1, cutoff=0.45)
    if not matches:
        return None
    idx = stems.index(matches[0])
    return files[idx]


def _resolve_source_path(raw_path, material_dir=DEFAULT_MATERIAL_DIR):
    raw_path = str(raw_path or "").strip()
    if not raw_path:
        return None
    expanded = os.path.expanduser(raw_path)
    candidates = []
    path = Path(expanded)
    if path.is_absolute():
        candidates.append(path)
    else:
        candidates.extend([
            Path(material_dir) / path,
            Path(PROJECT_DIR) / path,
            Path.cwd() / path,
        ])
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate.resolve()
    return None


def load_source_manifest(manifest_file=None, material_dir=DEFAULT_MATERIAL_DIR):
    """Load optional project -> original source file mapping.

    The expected local CSV columns are project_name and user_supplied_path.
    JSON is also accepted as either {project_name: path} or a list of records.
    Missing files are ignored so a partially filled manifest can be reused.
    """
    manifest_file = str(manifest_file or "").strip()
    if not manifest_file:
        return {}
    path = Path(os.path.expanduser(manifest_file))
    if not path.is_absolute():
        path = Path(PROJECT_DIR) / path
    if not path.exists():
        return {}

    mappings = {}

    def add_mapping(project_name, source_path):
        project_name = _clean_text(project_name)
        resolved = _resolve_source_path(source_path, material_dir=material_dir)
        if project_name and resolved:
            mappings[project_name] = resolved

    if path.suffix.lower() == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(payload, dict):
            for project_name, source_path in payload.items():
                if isinstance(source_path, dict):
                    source_path = (
                        source_path.get("user_supplied_path")
                        or source_path.get("source_path")
                        or source_path.get("path")
                    )
                add_mapping(project_name, source_path)
        elif isinstance(payload, list):
            for row in payload:
                if not isinstance(row, dict):
                    continue
                add_mapping(
                    row.get("project_name"),
                    row.get("user_supplied_path") or row.get("source_path") or row.get("path") or row.get("file_path"),
                )
        return mappings

    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            add_mapping(
                row.get("project_name"),
                row.get("user_supplied_path") or row.get("source_path") or row.get("path") or row.get("file_path"),
            )
    return mappings


_MATERIAL_TEXT_CACHE = {}


def _material_text_lines(path):
    path = Path(path)
    cache_key = str(path)
    if cache_key in _MATERIAL_TEXT_CACHE:
        return _MATERIAL_TEXT_CACHE[cache_key]

    lines = []
    suffix = path.suffix.lower()
    try:
        if suffix == ".xlsx":
            from openpyxl import load_workbook

            wb = load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.worksheets:
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    row_text = " | ".join(str(v).strip() for v in row if v is not None and str(v).strip())
                    if row_text:
                        lines.append({
                            "location": f"{sheet.title} 第{row_idx}行",
                            "text": row_text,
                        })
        elif suffix == ".docx":
            from docx import Document

            doc = Document(path)
            for idx, paragraph in enumerate(doc.paragraphs, start=1):
                text = paragraph.text.strip()
                if text:
                    lines.append({"location": f"段落{idx}", "text": text})
            for table_idx, table in enumerate(doc.tables, start=1):
                for row_idx, row in enumerate(table.rows, start=1):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        lines.append({"location": f"表{table_idx} 第{row_idx}行", "text": row_text})
    except Exception as exc:
        lines.append({"location": "解析失败", "text": str(exc)})

    _MATERIAL_TEXT_CACHE[cache_key] = lines
    return lines


def _material_full_text(path):
    return "\n".join(line.get("text", "") for line in _material_text_lines(path))


def _evidence_keywords(opinion, work_category=""):
    generic = {
        "广州", "万科", "工程", "施工", "方案", "报价", "清单", "明确", "需要", "缺失",
        "是否", "为什么", "检查", "验收", "项目", "修复", "翻新", "改造",
    }
    keywords = []
    for keyword in trigger_keywords_for(opinion, work_category=work_category):
        if keyword and keyword not in generic and len(keyword) >= 2:
            keywords.append(keyword)
    for checkpoint in expected_checkpoints_for(opinion, work_category):
        for alias in checkpoint["aliases"]:
            if alias not in generic and alias not in keywords:
                keywords.append(alias)
    for token in re.findall(r"[A-Za-z0-9]{2,}|[\u4e00-\u9fff]{2,8}", str(opinion or "")):
        if token not in generic and token not in keywords:
            keywords.append(token)
    return keywords[:12]


def _is_evidence_noise(line_text, opinion):
    text = str(line_text or "")
    opinion_text = str(opinion or "")
    if "保修" in text and "保修" not in opinion_text:
        return True
    if "屋面防水工程、有防水要求" in text and "保修" not in opinion_text:
        return True
    if "我司施工内容" in text and "施工内容" not in opinion_text and "界面" not in opinion_text:
        return True
    if "移交状态" in text and "移交" not in opinion_text:
        return True
    return False


def extract_scheme_evidence(row, material_dir=DEFAULT_MATERIAL_DIR, max_snippets=3):
    matched_file = row.get("matched_file")
    matched_path = row.get("matched_file_path")
    if not matched_file and not matched_path:
        return []
    path = Path(matched_path) if matched_path else Path(material_dir) / matched_file
    if not path.exists():
        return []
    lines = _material_text_lines(path)
    keywords = _evidence_keywords(row.get("opinion", ""), row.get("work_category", ""))
    scored = []
    for line in lines:
        text = line["text"]
        if _is_evidence_noise(text, row.get("opinion", "")):
            continue
        score = sum(1 for keyword in keywords if keyword.lower() in text.lower())
        if score:
            scored.append((score, len(text), line))
    scored.sort(key=lambda item: (-item[0], item[1]))
    evidence = []
    for score, _, line in scored[:max_snippets]:
        evidence.append({
            "source_file": matched_file,
            "location": line["location"],
            "matched_keywords": [kw for kw in keywords if kw.lower() in line["text"].lower()][:6],
            "text": _snippet_around_keywords(line["text"], keywords, max_chars=600),
            "match_score": score,
        })
    return evidence


def _matched_path_for_row(row, material_dir=DEFAULT_MATERIAL_DIR):
    matched_path = row.get("matched_file_path")
    matched_file = row.get("matched_file")
    if matched_path:
        return Path(matched_path)
    if matched_file:
        return Path(material_dir) / matched_file
    return None


def enrich_rows_with_scheme_evidence(rows, material_dir=DEFAULT_MATERIAL_DIR, max_snippets=3):
    enriched = []
    for row in rows:
        copy = dict(row)
        copy["scheme_evidence"] = extract_scheme_evidence(copy, material_dir=material_dir, max_snippets=max_snippets)
        if not copy["scheme_evidence"] and copy.get("source_match_quality_label") in {"高", "中", "人工映射"}:
            matched_path = _matched_path_for_row(copy, material_dir=material_dir)
            if matched_path and matched_path.exists():
                copy["alignment_text"] = _material_full_text(matched_path)
        copy.update(assess_scheme_alignment(copy))
        copy.pop("alignment_text", None)
        enriched.append(copy)
    return enriched


def load_opinion_rows(opinion_file=DEFAULT_OPINION_FILE, material_dir=DEFAULT_MATERIAL_DIR, source_manifest=None):
    from openpyxl import load_workbook

    wb = load_workbook(opinion_file, data_only=True)
    ws = wb.active
    files = index_material_files(material_dir)
    source_map = load_source_manifest(source_manifest, material_dir=material_dir)
    rows = []
    for row_index, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        project_name, engineer, opinion = row[:3]
        if not project_name:
            continue
        project_name = _clean_text(project_name)
        manifest_file = source_map.get(project_name)
        material_file = manifest_file or match_material_file(project_name, files)
        source_match_type = "manifest" if manifest_file else ("fuzzy" if material_file else "")
        if manifest_file:
            match_quality = {"label": "人工映射", "score": 1.0}
        else:
            match_quality = material_match_quality(project_name, material_file)
        project_type = classify_project_type(project_name)
        for item_index, item in enumerate(split_opinion_items(opinion), start=1):
            work_category = infer_work_category(f"{project_name} {item}")
            dimension = classify_dimension(item)
            evidence_type = infer_evidence_type(item, dimension)
            attribution = build_attribution(item, dimension, work_category)
            rows.append({
                "row_index": row_index,
                "item_index": item_index,
                "project_name": project_name,
                "engineer": _clean_text(engineer),
                "opinion": _clean_text(item),
                "project_type": project_type,
                "matched_file": material_file.name if material_file else "",
                "matched_file_path": str(material_file) if material_file else "",
                "source_match_type": source_match_type,
                "source_match_quality_label": match_quality["label"],
                "source_match_quality_score": match_quality["score"],
                "file_type": material_file.suffix.lower().lstrip(".") if material_file else "",
                "work_category": work_category,
                "dimension": dimension,
                "is_scheme_related": project_type in {"施工方案", "方案兼报价"} or "方案" in project_name,
                "is_cost_related": project_type in {"报价/白单", "方案兼报价"},
                "evidence_type": evidence_type,
                "evidence_ref": infer_evidence_ref(item, work_category, evidence_type),
                **attribution,
            })
    return rows


def should_promote_to_experience(row, scope="scheme-priority"):
    opinion = _clean_text(row.get("opinion", ""))
    if is_non_actionable_opinion(opinion):
        return False
    if len(opinion) <= 8 and re.search(r"[:：]$", opinion):
        return False
    if scope == "all":
        return True
    if scope == "scheme-only":
        return bool(row.get("is_scheme_related"))
    if row.get("is_scheme_related"):
        return True
    return row.get("dimension") in {"描述完整性", "分项拆分", "逻辑自洽"} and bool(row.get("opinion"))


def opinion_row_to_card(row):
    opinion = row["opinion"]
    dimension = row["dimension"]
    work_category = row["work_category"]
    extensions = extension_rules_for(opinion, row["project_name"])
    trigger_keywords = trigger_keywords_for(opinion, row["project_name"], work_category)
    confidence = "高" if row.get("is_scheme_related") and trigger_keywords else "中"
    attribution = build_attribution(opinion, dimension, work_category)
    alignment = (
        {
            "alignment_status": row.get("alignment_status"),
            "covered_points": row.get("covered_points", []),
            "partial_points": row.get("partial_points", []),
            "missing_points": row.get("missing_points", []),
            "checkpoint_assessments": row.get("checkpoint_assessments", []),
            "scheme_gap": row.get("scheme_gap"),
            "expert_intent": row.get("expert_intent"),
            "evidence_chain": row.get("evidence_chain"),
        }
        if row.get("alignment_status")
        else assess_scheme_alignment({**row, **attribution})
    )
    return {
        "id": _hash_id("EXP", row.get("project_name"), row.get("row_index"), row.get("item_index"), opinion),
        "source_project": row.get("project_name", ""),
        "source_engineer": row.get("engineer", ""),
        "source_row": row.get("row_index", -1),
        "source_item": row.get("item_index", -1),
        "source_opinion": opinion,
        "project_type": row.get("project_type", ""),
        "file_type": row.get("file_type", ""),
        "matched_file": row.get("matched_file", ""),
        "work_category": work_category,
        "dimension": dimension,
        "background": build_background(opinion, dimension, work_category),
        "reason": build_reason(opinion, dimension),
        "problem_pattern": row.get("problem_pattern", attribution["problem_pattern"]),
        "problem_pattern_label": row.get("problem_pattern_label", attribution["problem_pattern_label"]),
        "secondary_patterns": row.get("secondary_patterns", attribution["secondary_patterns"]),
        "professional_attribution": row.get("professional_attribution", attribution["professional_attribution"]),
        "professional_attribution_label": row.get("professional_attribution_label", attribution["professional_attribution_label"]),
        "secondary_professional_attributions": row.get(
            "secondary_professional_attributions",
            attribution["secondary_professional_attributions"],
        ),
        "engineer_question": row.get("engineer_question", attribution["engineer_question"]),
        "review_intents": row.get("review_intents", attribution["review_intents"]),
        "root_cause": row.get("root_cause", attribution["root_cause"]),
        "risk_if_ignored": row.get("risk_if_ignored", attribution["risk_if_ignored"]),
        "generalization_rule": row.get("generalization_rule", attribution["generalization_rule"]),
        "review_questions": row.get("review_questions", attribution["review_questions"]),
        "required_artifacts": row.get("required_artifacts", attribution["required_artifacts"]),
        "fix_template": row.get("fix_template", attribution["fix_template"]),
        "standard_query": row.get("standard_query", attribution["standard_query"]),
        "scheme_evidence": row.get("scheme_evidence", []),
        "alignment_status": alignment.get("alignment_status", "无法判断"),
        "covered_points": alignment.get("covered_points", []),
        "partial_points": alignment.get("partial_points", []),
        "missing_points": alignment.get("missing_points", []),
        "checkpoint_assessments": alignment.get("checkpoint_assessments", []),
        "scheme_gap": alignment.get("scheme_gap", ""),
        "expert_intent": alignment.get("expert_intent", ""),
        "evidence_chain": alignment.get("evidence_chain", {}),
        "evidence_type": row.get("evidence_type", "专家经验"),
        "evidence_ref": row.get("evidence_ref", "历史审核经验：零星工程专家意见"),
        "extension_rules": extensions,
        "trigger_keywords": trigger_keywords,
        "confidence": confidence,
        "created_at": _dt.datetime.now().isoformat(timespec="seconds"),
    }


def build_experience_cards(rows, scope="scheme-priority"):
    return [opinion_row_to_card(row) for row in rows if should_promote_to_experience(row, scope=scope)]


def card_to_kb_rule(card):
    extension_text = "；".join(rule["rule"] for rule in card.get("extension_rules", []))
    evidence_text = "；".join(
        f"{item.get('location')}：{item.get('text')}"
        for item in card.get("scheme_evidence", [])[:2]
    )
    content = (
        f"【零星工程审核经验】\n"
        f"适用类别：{card.get('work_category')}\n"
        f"审核维度：{card.get('dimension')}\n"
        f"触发关键词：{'、'.join(card.get('trigger_keywords', []))}\n"
        f"历史意见：{card.get('source_opinion')}\n"
        f"问题背景：{card.get('background')}\n"
        f"原因：{card.get('reason')}\n"
        f"问题模式：{card.get('problem_pattern_label')}({card.get('problem_pattern')})\n"
        f"专业归因：{card.get('professional_attribution_label')}({card.get('professional_attribution')})\n"
        f"工程师实际问题：{card.get('engineer_question')}\n"
        f"审核意图：{'、'.join(card.get('review_intents', []))}\n"
        f"根因归因：{card.get('root_cause')}\n"
        f"忽略风险：{card.get('risk_if_ignored')}\n"
        f"泛化规则：{card.get('generalization_rule')}\n"
        f"证据对齐状态：{card.get('alignment_status')}\n"
        f"已覆盖控制点：{'、'.join(card.get('covered_points', []))}\n"
        f"笼统提及控制点：{'、'.join(card.get('partial_points', []))}\n"
        f"缺失控制点：{'、'.join(card.get('missing_points', []))}\n"
        f"方案缺口判断：{card.get('scheme_gap')}\n"
        f"专家真实意图：{card.get('expert_intent')}\n"
        f"复核问题：{'；'.join(card.get('review_questions', []))}\n"
        f"应补资料：{'、'.join(card.get('required_artifacts', []))}\n"
        f"方案证据：{evidence_text}\n"
        f"修改模板：{card.get('fix_template')}\n"
        f"延伸检查：{extension_text}\n"
        f"依据类型：{card.get('evidence_type')}\n"
        f"依据出处：{card.get('evidence_ref')}\n"
    )
    return {
        "id": card["id"],
        "category": "零星改造审核经验",
        "wbs_code": "通用",
        "level": 2,
        "content": content,
        "tags": card.get("trigger_keywords", []) + [card.get("dimension", ""), card.get("work_category", "")],
        "is_washed": True,
        "condensed_content": (
            f"{card.get('work_category')}｜{card.get('dimension')}｜"
            f"{card.get('professional_attribution_label')}｜{card.get('alignment_status')}｜"
            f"{card.get('source_opinion')}｜缺口：{card.get('scheme_gap')}｜泛化：{card.get('generalization_rule')}"
        )[:700],
        "ingest_time": card.get("created_at", ""),
        "source_file": "零星工程历史审核经验",
        "seq_index": int(card.get("source_row", -1)) * 100 + int(card.get("source_item", 0)),
        "status": "active",
        "full_text": json.dumps(card, ensure_ascii=False),
        "summary": card.get("reason", ""),
        "publish_date": "2026-04-30",
        "lifecycle_phase": "施工",
        "index_source": "review_experience",
        "node_id": card["id"],
        "node_title": f"{card.get('work_category')}｜{card.get('dimension')}",
        "quality_score": 95,
        "quality_flags": [],
        "quality_notes": "从历史审核意见结构化抽取的专家经验卡",
    }


def summarize_rows(rows):
    return {
        "total_items": len(rows),
        "projects": len({row.get("project_name", "") for row in rows if row.get("project_name")}),
        "engineers": Counter(row.get("engineer", "") for row in rows),
        "dimensions": Counter(row.get("dimension", "描述完整性") for row in rows),
        "work_categories": Counter(row.get("work_category", "通用零星维修") for row in rows),
        "project_types": Counter(row.get("project_type", "未知") for row in rows),
        "problem_patterns": Counter(row.get("problem_pattern", "expert_judgement") for row in rows),
        "professional_attributions": Counter(
            row.get("professional_attribution", "expert_engineering_judgement") for row in rows
        ),
        "alignment_statuses": Counter(row.get("alignment_status", "无法判断") for row in rows),
        "review_intents": Counter(intent for row in rows for intent in row.get("review_intents", [])),
    }


def build_methodology(rows, cards):
    summary = summarize_rows(rows)
    pattern_examples = defaultdict(list)
    professional_examples = defaultdict(list)
    for row in rows:
        if len(pattern_examples[row.get("problem_pattern", "expert_judgement")]) < 5:
            pattern_examples[row.get("problem_pattern", "expert_judgement")].append({
                "project_name": row.get("project_name", ""),
                "opinion": row.get("opinion", ""),
                "work_category": row.get("work_category", "通用零星维修"),
                "dimension": row.get("dimension", "描述完整性"),
                "root_cause": row.get("root_cause", ""),
                "generalization_rule": row.get("generalization_rule", ""),
                "alignment_status": row.get("alignment_status", "无法判断"),
                "scheme_gap": row.get("scheme_gap", ""),
                "partial_points": row.get("partial_points", []),
            })
        professional_code = row.get("professional_attribution", "expert_engineering_judgement")
        if len(professional_examples[professional_code]) < 6:
            professional_examples[professional_code].append({
                "project_name": row.get("project_name", ""),
                "opinion": row.get("opinion", ""),
                "work_category": row.get("work_category", "通用零星维修"),
                "engineer_question": row.get("engineer_question", ""),
                "scheme_evidence": row.get("scheme_evidence", [])[:2],
                "alignment_status": row.get("alignment_status", "无法判断"),
                "scheme_gap": row.get("scheme_gap", ""),
                "partial_points": row.get("partial_points", []),
            })

    patterns = []
    for code, count in summary["problem_patterns"].most_common():
        label = PATTERN_LABELS.get(code, "专家经验判断")
        sample = pattern_examples.get(code, [])
        root_cause = sample[0]["root_cause"] if sample else ""
        generalization_rule = sample[0]["generalization_rule"] if sample else ""
        patterns.append({
            "code": code,
            "label": label,
            "count": count,
            "root_cause": root_cause,
            "generalization_rule": generalization_rule,
            "examples": sample,
        })

    category_methods = []
    for category, _ in summary["work_categories"].most_common():
        category_rows = [row for row in rows if row.get("work_category", "通用零星维修") == category]
        category_methods.append({
            "work_category": category,
            "count": len(category_rows),
            "top_patterns": Counter(row.get("problem_pattern", "expert_judgement") for row in category_rows).most_common(5),
            "top_dimensions": Counter(row.get("dimension", "描述完整性") for row in category_rows).most_common(4),
            "focus": category_focus(category),
        })

    return {
        "version": "2026-04-30.v2_repair_methodology",
        "core_purpose": "判断班组方案能否指导施工、计价、验收和复核，而不是做大而全的施工组织设计审查。",
        "review_loop": [
            "拆分分项工程 WorkItem",
            "识别现场条件、材料参数、工序顺序、验收指标和清单口径",
            "按四个核心维度检查可施工、可计价、可验收、可复核",
            "用规范库验证硬性要求，用历史经验解释专家关注点",
            "输出问题、背景原因、依据类型、依据出处和可直接修改的建议",
        ],
        "problem_patterns": patterns,
        "professional_attributions": [
            {
                "code": code,
                "label": PROFESSIONAL_LABELS.get(code, "需专业工程师结合上下文判断"),
                "count": count,
                "examples": professional_examples.get(code, []),
            }
            for code, count in summary["professional_attributions"].most_common()
        ],
        "category_methods": category_methods,
        "review_intents": dict(summary["review_intents"]),
        "alignment_statuses": dict(summary["alignment_statuses"]),
        "experience_card_count": len(cards),
    }


def category_focus(category):
    focus = {
        "防水渗漏": ["部位拆分", "材料适用性", "基层处理", "收口节点", "闭水/淋水验证"],
        "地坪/EPDM/环氧": ["基层验收", "材料配比", "各层厚度", "固化养护", "相邻水沟/边界交接"],
        "墙面涂料": ["旧基层处理", "腻子/底漆/面漆遍数", "材料品类", "打磨厚度", "观感一致性"],
        "瓷砖铺贴": ["瓷砖规格", "胶粘材料", "铺贴厚度", "空鼓检查", "标高和污染控制"],
        "门窗玻璃": ["防火/安全认证", "五金配件", "厚度规格", "安装节点", "型式检验/3C"],
        "给排水管网": ["排水问题定义", "管材和井盖等级", "沟槽/管井做法", "通水试验", "道路承载"],
        "弱电设备": ["设备参数", "传输距离", "供电和线缆", "存储天数", "调试试运行"],
        "结构修补": ["植筋锚固", "孔位错开", "结构风险", "隐蔽验收", "替代结构方案"],
        "装修翻新": ["拆除恢复边界", "基层条件", "材料做法", "验收指标", "与白单一致性"],
    }
    return focus.get(category, ["材料参数", "施工做法", "工序顺序", "验收指标", "清单口径"])


REPRESENTATIVE_CASE_SPECS = [
    {
        "case_label": "地坪/EPDM样本",
        "hints": ("EPDM", "塑胶", "水沟", "固化"),
        "expected_keywords": ["EPDM胶水比", "固化时间", "基层验收", "水沟", "成品保护"],
    },
    {
        "case_label": "宿舍/结构改造样本",
        "hints": ("反坎", "抹灰", "植筋", "错孔", "结构隔层"),
        "expected_keywords": ["200mm混凝土反坎", "轻质砂浆", "植筋深度", "错位"],
    },
    {
        "case_label": "户外/电梯/活动室综合改造样本",
        "hints": ("角铁", "方通", "瓷砖胶", "防护剂", "3C", "1底1面", "电梯", "大理石"),
        "expected_keywords": ["角铁尺寸", "方通", "瓷砖胶", "防护剂", "3C", "1底1面"],
    },
]


def _project_text(rows):
    return " ".join(
        str(row.get(key, ""))
        for row in rows
        for key in ("project_name", "opinion", "work_category", "dimension")
    )


def select_representative_projects(rows, specs=None):
    specs = specs or REPRESENTATIVE_CASE_SPECS
    grouped = defaultdict(list)
    for row in rows:
        grouped[row.get("project_name", "")].append(row)

    selected = []
    used_projects = set()
    for spec in specs:
        best_project = ""
        best_score = 0
        for project_name, project_rows in grouped.items():
            if not project_name or project_name in used_projects:
                continue
            text = _project_text(project_rows)
            score = sum(text.count(hint) for hint in spec["hints"])
            if score > best_score:
                best_project = project_name
                best_score = score
        if best_project:
            used_projects.add(best_project)
            selected.append((best_project, spec))
    return selected


def build_benchmark_cases(rows):
    cases = []
    for target, spec in select_representative_projects(rows):
        expected = spec["expected_keywords"]
        related = [row for row in rows if row["project_name"] == target]
        cases.append({
            "case_label": spec["case_label"],
            "project_name": target,
            "expected_keywords": expected,
            "source_opinions": [row["opinion"] for row in related],
            "dimensions": sorted({row["dimension"] for row in related}),
            "work_categories": sorted({row["work_category"] for row in related}),
        })
    return cases


def build_deep_attribution_cases(rows, max_per_category=8):
    grouped = defaultdict(list)
    for row in rows:
        code = row.get("professional_attribution", "expert_engineering_judgement")
        if len(grouped[code]) >= max_per_category:
            continue
        grouped[code].append({
            "project_name": row.get("project_name", ""),
            "opinion": row.get("opinion", ""),
            "work_category": row.get("work_category", ""),
            "dimension": row.get("dimension", ""),
            "professional_attribution": code,
            "professional_attribution_label": row.get("professional_attribution_label", PROFESSIONAL_LABELS.get(code, "")),
            "engineer_question": row.get("engineer_question", ""),
            "root_cause": row.get("root_cause", ""),
            "generalization_rule": row.get("generalization_rule", ""),
            "review_questions": row.get("review_questions", []),
            "required_artifacts": row.get("required_artifacts", []),
            "scheme_evidence": row.get("scheme_evidence", []),
            "alignment_status": row.get("alignment_status", "无法判断"),
            "covered_points": row.get("covered_points", []),
            "partial_points": row.get("partial_points", []),
            "missing_points": row.get("missing_points", []),
            "checkpoint_assessments": row.get("checkpoint_assessments", []),
            "scheme_gap": row.get("scheme_gap", ""),
            "expert_intent": row.get("expert_intent", ""),
            "evidence_chain": row.get("evidence_chain", {}),
        })
    return {
        "version": "2026-04-30.deep_attribution",
        "note": "Evidence snippets are extracted from matched local material files when a keyword overlap is found.",
        "categories": [
            {
                "code": code,
                "label": PROFESSIONAL_LABELS.get(code, "需专业工程师结合上下文判断"),
                "cases": cases,
            }
            for code, cases in grouped.items()
        ],
    }


def build_alignment_cases(rows):
    cases = []
    for target, spec in select_representative_projects(rows):
        related = [row for row in rows if row.get("project_name") == target]
        items = []
        for row in related:
            alignment = row if row.get("alignment_status") else {**row, **assess_scheme_alignment(row)}
            items.append({
                "source_item": row.get("item_index"),
                "opinion": row.get("opinion", ""),
                "dimension": row.get("dimension", ""),
                "work_category": row.get("work_category", ""),
                "professional_attribution_label": row.get("professional_attribution_label", ""),
                "expert_intent": alignment.get("expert_intent", ""),
                "alignment_status": alignment.get("alignment_status", "无法判断"),
                "covered_points": alignment.get("covered_points", []),
                "partial_points": alignment.get("partial_points", []),
                "missing_points": alignment.get("missing_points", []),
                "checkpoint_assessments": alignment.get("checkpoint_assessments", []),
                "scheme_gap": alignment.get("scheme_gap", ""),
                "scheme_evidence": row.get("scheme_evidence", [])[:3],
                "review_questions": row.get("review_questions", []),
                "required_artifacts": row.get("required_artifacts", []),
                "generalization_rule": row.get("generalization_rule", ""),
                "risk_if_ignored": row.get("risk_if_ignored", ""),
            })
        cases.append({
            "case_label": spec["case_label"],
            "project_name": target,
            "items": items,
            "status_distribution": dict(Counter(item["alignment_status"] for item in items)),
        })
    return {
        "version": "2026-04-30.deep_alignment",
        "method": "逐条审核意见对照原方案证据，记录已覆盖/缺失控制点，再提炼专家追问和可迁移规则。",
        "cases": cases,
    }


def render_deep_alignment_report(alignment_cases):
    lines = [
        "# 原始方案-审核意见深度对照报告",
        "",
        "## 阅读方式",
        "- 每条意见都按 `原意见 -> 原方案证据 -> 对齐状态 -> 专家真实追问 -> 缺口原因 -> 泛化规则` 展开。",
        "- `已补齐` 不代表历史意见无价值，而是说明当前留存方案可能已经按意见修订；系统运行时不能机械复述。",
        "- `部分补齐/仍缺失` 才是更适合转化为运行期检查清单的经验。",
        "",
    ]
    for case in alignment_cases.get("cases", []):
        lines += [
            f"## {case.get('project_name')}",
            f"- 对齐状态分布：{json.dumps(case.get('status_distribution', {}), ensure_ascii=False)}",
            "",
        ]
        for item in case.get("items", []):
            lines += [
                f"### {item.get('source_item')}. {item.get('opinion')}",
                f"- 维度/类别：{item.get('dimension')}｜{item.get('work_category')}｜{item.get('professional_attribution_label')}",
                f"- 对齐状态：{item.get('alignment_status')}",
                f"- 具体覆盖控制点：{'、'.join(item.get('covered_points', [])) or '未识别'}",
                f"- 笼统提及控制点：{'、'.join(item.get('partial_points', [])) or '未识别'}",
                f"- 缺失控制点：{'、'.join(item.get('missing_points', [])) or '未识别'}",
                f"- 专家真实意图：{item.get('expert_intent')}",
                f"- 缺口原因：{item.get('scheme_gap')}",
                f"- 忽略风险：{item.get('risk_if_ignored')}",
                f"- 泛化规则：{item.get('generalization_rule')}",
                "- 控制点判断：",
            ]
            assessments = item.get("checkpoint_assessments", [])
            if assessments:
                for assessment in assessments:
                    lines.append(
                        f"  - {assessment.get('name')}：{assessment.get('status')}，{assessment.get('note')}"
                    )
            else:
                lines.append("  - 未形成可拆解控制点")
            lines += [
                "- 原方案证据：",
            ]
            evidence = item.get("scheme_evidence", [])
            if evidence:
                for snippet in evidence:
                    lines.append(f"  - {snippet.get('location')}：{snippet.get('text')}")
            else:
                lines.append("  - 未匹配到可对照片段")
            lines.append("")
    return "\n".join(lines)


def write_analysis_outputs(rows, cards, output_dir=ANALYSIS_DIR):
    os.makedirs(output_dir, exist_ok=True)
    summary = summarize_rows(rows)
    benchmark_cases = build_benchmark_cases(rows)
    methodology = build_methodology(rows, cards)
    deep_cases = build_deep_attribution_cases(rows)
    alignment_cases = build_alignment_cases(rows)
    cards_path = os.path.join(output_dir, "review_experience_cards.json")
    benchmark_path = os.path.join(output_dir, "review_benchmark_cases.json")
    methodology_path = os.path.join(output_dir, "review_methodology.json")
    deep_cases_path = os.path.join(output_dir, "review_deep_attribution_cases.json")
    alignment_cases_path = os.path.join(output_dir, "review_alignment_cases.json")
    alignment_report_path = os.path.join(output_dir, "deep_alignment_benchmark_report.md")
    report_path = os.path.join(output_dir, "raw_material_review_report.md")

    with open(cards_path, "w", encoding="utf-8") as f:
        json.dump(cards, f, ensure_ascii=False, indent=2)
    with open(benchmark_path, "w", encoding="utf-8") as f:
        json.dump(benchmark_cases, f, ensure_ascii=False, indent=2)
    with open(methodology_path, "w", encoding="utf-8") as f:
        json.dump(methodology, f, ensure_ascii=False, indent=2)
    with open(deep_cases_path, "w", encoding="utf-8") as f:
        json.dump(deep_cases, f, ensure_ascii=False, indent=2)
    with open(alignment_cases_path, "w", encoding="utf-8") as f:
        json.dump(alignment_cases, f, ensure_ascii=False, indent=2)
    with open(alignment_report_path, "w", encoding="utf-8") as f:
        f.write(render_deep_alignment_report(alignment_cases))
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(render_analysis_report(rows, cards, summary, benchmark_cases, methodology))
    return {
        "report_path": report_path,
        "cards_path": cards_path,
        "benchmark_path": benchmark_path,
        "methodology_path": methodology_path,
        "deep_cases_path": deep_cases_path,
        "alignment_cases_path": alignment_cases_path,
        "alignment_report_path": alignment_report_path,
    }


def _counter_lines(counter):
    return "\n".join(f"- {key}: {value}" for key, value in counter.most_common())


def render_analysis_report(rows, cards, summary, benchmark_cases, methodology=None):
    examples = defaultdict(list)
    for row in rows:
        if len(examples[row["dimension"]]) < 6:
            examples[row["dimension"]].append(row)

    lines = [
        "# 原始材料审核方法论分析报告",
        "",
        "## 核心结论",
        "- 这些材料面向零星维修/改造，不是大体量施工组织设计。",
        "- 人类审核重点是方案能否指导班组施工、报价复核、过程验收和结算界面。",
        "- 高频问题集中在描述完整性、工艺合理性、分项拆分和逻辑自洽。",
        "- 历史审核意见应作为专家经验逐条结构化，不能整表作为一条知识库规则。",
        "",
        "## 样本统计",
        f"- 原子审核意见: {summary['total_items']}",
        f"- 涉及项目: {summary['projects']}",
        f"- 结构化经验卡: {len(cards)}",
        "",
        "### 审核维度分布",
        _counter_lines(summary["dimensions"]),
        "",
        "### 工程类别分布",
        _counter_lines(summary["work_categories"]),
        "",
        "### 资料类型分布",
        _counter_lines(summary["project_types"]),
        "",
        "### 问题模式分布",
        _counter_lines(Counter({PATTERN_LABELS.get(k, "专家经验判断"): v for k, v in summary["problem_patterns"].items()})),
        "",
        "### 专业归因分布",
        _counter_lines(Counter({
            PROFESSIONAL_LABELS.get(k, "需专业工程师结合上下文判断"): v
            for k, v in summary["professional_attributions"].items()
        })),
        "",
        "### 原方案证据对齐状态",
        _counter_lines(summary["alignment_statuses"]),
        "",
        "### 审核意图分布",
        _counter_lines(summary["review_intents"]),
        "",
        "## 四类核心审核目的",
        "1. 描述完整性：材料、规格、厚度、强度、配比、范围、验收指标必须写到能采购、能施工、能复核。",
        "2. 工艺合理性：做法要适配基层、场景和耐久性，避免班组套模板或选错材料。",
        "3. 分项拆分：方案、白单/清单、施工动作和结算口径要一致，避免漏项、重复计价和责任不清。",
        "4. 逻辑自洽：工序顺序、参数、现场条件、前后章节和验收动作不能互相打架。",
        "",
        "## 归因方法论",
        "- 先看意见是在补参数、纠工艺、拆计价、查逻辑、问现场、补验收还是清界面。",
        "- 再反推专家的真正目的：让方案能指导施工、支撑计价、便于验收和留下复核证据。",
        "- 最后把意见泛化成同类分项的检查动作，而不是照搬某个项目的原句。",
    ]
    if methodology:
        lines += ["", "## 专业工程师归因框架"]
        for item in methodology.get("professional_attributions", [])[:12]:
            lines += ["", f"### {item['label']}", f"- 样本数：{item['count']}"]
            for example in item.get("examples", [])[:3]:
                evidence = example.get("scheme_evidence", [])
                evidence_text = f"｜方案证据：{evidence[0]['location']} {evidence[0]['text'][:120]}" if evidence else ""
                lines.append(
                    f"- {example['project_name']}｜{example['opinion']}｜"
                    f"状态：{example.get('alignment_status', '无法判断')}｜"
                    f"工程师追问：{example.get('engineer_question', '')}｜"
                    f"缺口：{example.get('scheme_gap', '')}{evidence_text}"
                )
        for pattern in methodology.get("problem_patterns", [])[:8]:
            lines += [
                "",
                f"### {pattern['label']}",
                f"- 归因：{pattern.get('root_cause', '')}",
                f"- 泛化：{pattern.get('generalization_rule', '')}",
            ]
            for example in pattern.get("examples", [])[:3]:
                lines.append(f"- 样本：{example['project_name']}｜{example['opinion']}")
        lines += ["", "## 分项工程审核焦点"]
        for item in methodology.get("category_methods", [])[:10]:
            lines.append(f"- {item['work_category']}：{'、'.join(item.get('focus', []))}")
    lines += [
        "",
        "## 典型意见样本",
    ]
    for dimension in CORE_DIMENSIONS:
        lines += ["", f"### {dimension}"]
        for row in examples.get(dimension, []):
            lines.append(f"- {row['project_name']}：{row['opinion']}")
    lines += [
        "",
        "## 当前系统偏差",
        "- 旧链路按 Agent 维度输出，安全、合同、造价容易喧宾夺主。",
        "- 方案型 Excel 被粗切成整张表，LLM 很难定位具体分项。",
        "- `城市公司检查结果` 整表曾作为单条 active 规则入库，检索命中后会把大量无关历史意见一次性喂给模型。",
        "- 审核结果应按分项工程组织，而不是按 Agent 名称组织。",
        "",
        "## 基准案例",
    ]
    for case in benchmark_cases:
        lines.append(f"- {case['project_name']}：期望命中 {'、'.join(case['expected_keywords'])}")
    return "\n".join(lines) + "\n"


def load_analysis_cards(path=None):
    path = path or os.path.join(ANALYSIS_DIR, "review_experience_cards.json")
    if not os.path.exists(path):
        return []
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data if isinstance(data, list) else []


def match_experience_cards(text, cards, limit=8, min_overlap=2):
    haystack = str(text or "").lower()
    generic_keywords = {
        "广州", "万科", "工程", "施工", "方案", "报价", "清单", "翻新", "维修", "改造",
        "明确", "缺失", "需要", "项目", "修缮", "进行", "材料", "工艺",
    }
    strong_keywords = {
        "EPDM", "环氧", "水沟", "轻质砖", "反坎", "植筋", "聚氨酯", "瓷砖",
        "钢化玻璃", "大理石", "防火门", "管井", "摄像", "人脸识别", "脚手架",
        "方通", "角铁", "C2TE", "3C",
    }
    scored = []
    for card in cards:
        keywords = [
            str(kw).strip()
            for kw in card.get("trigger_keywords", [])
            if str(kw).strip() and str(kw).strip() not in generic_keywords
        ]
        overlap = sum(1 for keyword in keywords if str(keyword).lower() in haystack)
        has_strong_hit = any(keyword.lower() in haystack for keyword in strong_keywords if keyword in keywords)
        if overlap < min_overlap and not (min_overlap <= 2 and has_strong_hit):
            continue
        scored.append((overlap, card.get("confidence") == "高", card))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [item[2] for item in scored[:limit]]
